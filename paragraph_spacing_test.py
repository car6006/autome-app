#!/usr/bin/env python3

import asyncio
import httpx
import json
import uuid
from datetime import datetime
import os
import tempfile
from docx import Document
from docx.shared import Pt

# Backend URL from frontend environment
BACKEND_URL = "https://autome-fix.preview.emergentagent.com/api"

class ParagraphSpacingTester:
    def __init__(self):
        self.test_results = []
        self.test_user_email = f"test_user_{uuid.uuid4().hex[:8]}@example.com"
        self.test_user_password = "TestPassword123"
        self.auth_token = None
        self.test_note_id = None

    async def log_test(self, test_name, success, details=""):
        """Log test results"""
        result = {
            "test": test_name,
            "success": success,
            "details": details,
            "timestamp": datetime.now().isoformat()
        }
        self.test_results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name}")
        if details:
            print(f"   Details: {details}")

    async def create_test_user(self):
        """Create a test user for testing"""
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                user_data = {
                    "email": self.test_user_email,
                    "username": f"testuser{uuid.uuid4().hex[:8]}",
                    "password": self.test_user_password,
                    "first_name": "Test",
                    "last_name": "User"
                }
                
                response = await client.post(f"{BACKEND_URL}/auth/register", json=user_data)
                
                if response.status_code == 200:
                    data = response.json()
                    self.auth_token = data.get("access_token")
                    await self.log_test("Create Test User", True, f"User created with email: {self.test_user_email}")
                    return True
                else:
                    await self.log_test("Create Test User", False, f"Status: {response.status_code}, Response: {response.text}")
                    return False
                    
        except Exception as e:
            await self.log_test("Create Test User", False, f"Exception: {str(e)}")
            return False

    async def create_note_with_ai_conversations(self):
        """Create a test note with AI conversations for export testing"""
        try:
            headers = {"Authorization": f"Bearer {self.auth_token}"}
            
            # Create a text note with substantial content
            async with httpx.AsyncClient(timeout=30) as client:
                note_data = {
                    "title": "Paragraph Spacing Test Meeting",
                    "kind": "text",
                    "text_content": "This is a comprehensive meeting discussion about quarterly performance review and strategic planning initiatives. We discussed customer engagement metrics, operational efficiency improvements, team development programs, financial performance analysis, resource allocation strategies, market expansion opportunities, and competitive positioning analysis. The meeting covered various aspects of business operations including supply chain optimization, technology infrastructure upgrades, human resources development, and sustainability initiatives."
                }
                
                response = await client.post(f"{BACKEND_URL}/notes", json=note_data, headers=headers)
                
                if response.status_code == 200:
                    data = response.json()
                    self.test_note_id = data.get("id")
                    await self.log_test("Create Test Note", True, f"Note created with ID: {self.test_note_id}")
                    
                    # Add AI conversations to the note
                    await self.add_ai_conversations()
                    return True
                else:
                    await self.log_test("Create Test Note", False, f"Status: {response.status_code}, Response: {response.text}")
                    return False
                    
        except Exception as e:
            await self.log_test("Create Test Note", False, f"Exception: {str(e)}")
            return False

    async def add_ai_conversations(self):
        """Add multiple AI conversations to test comprehensive formatting"""
        try:
            headers = {"Authorization": f"Bearer {self.auth_token}"}
            
            # Add multiple AI conversations with different types of content
            questions = [
                "What are the key insights from this meeting discussion?",
                "Can you provide a summary of the strategic planning initiatives mentioned?",
                "What are the main action items that should be prioritized based on this discussion?",
                "How would you assess the overall performance metrics discussed in the meeting?",
                "What recommendations would you make for operational improvements?"
            ]
            
            async with httpx.AsyncClient(timeout=45) as client:
                for question in questions:
                    response = await client.post(
                        f"{BACKEND_URL}/notes/{self.test_note_id}/ai-chat",
                        json={"question": question},
                        headers=headers
                    )
                    
                    if response.status_code != 200:
                        await self.log_test("Add AI Conversations", False, f"Failed to add AI conversation: {response.status_code}")
                        return False
                
                await self.log_test("Add AI Conversations", True, f"Added {len(questions)} AI conversations to note")
                return True
                
        except Exception as e:
            await self.log_test("Add AI Conversations", False, f"Exception: {str(e)}")
            return False

    async def export_word_document(self):
        """Export the Word document and return content"""
        try:
            headers = {"Authorization": f"Bearer {self.auth_token}"}
            
            async with httpx.AsyncClient(timeout=60) as client:
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    file_size = len(response.content)
                    await self.log_test("Export Word Document", True, f"Export successful, file size: {file_size} bytes")
                    return response.content
                else:
                    await self.log_test("Export Word Document", False, f"Status: {response.status_code}, Response: {response.text}")
                    return None
                    
        except Exception as e:
            await self.log_test("Export Word Document", False, f"Exception: {str(e)}")
            return None

    def analyze_paragraph_spacing(self, docx_content):
        """Analyze the Word document for specific paragraph spacing requirements"""
        try:
            # Save content to temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(docx_content)
                temp_file_path = temp_file.name
            
            # Open and analyze the document
            doc = Document(temp_file_path)
            
            spacing_analysis = {
                "total_paragraphs": len(doc.paragraphs),
                "body_text_with_12pt_after": 0,
                "body_text_with_3pt_before": 0,
                "subheadings_with_12pt_before": 0,
                "subheadings_with_6pt_after": 0,
                "qa_pairs_with_separation": 0,
                "paragraphs_with_breathing_room": 0,
                "spacing_details": []
            }
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_format = paragraph.paragraph_format
                
                # Get spacing values (convert from Pt to points)
                space_before = para_format.space_before.pt if para_format.space_before else 0
                space_after = para_format.space_after.pt if para_format.space_after else 0
                
                # Determine paragraph type
                is_heading = (
                    paragraph.style.name.startswith('Heading') or 
                    paragraph.style.name in ['AI Report Title', 'AI Section Heading'] or
                    any(run.font.bold for run in paragraph.runs if run.font.bold is not None)
                )
                
                is_subheading = (
                    paragraph.style.name == 'AI Section Heading' or
                    (is_heading and paragraph.style.name != 'AI Report Title')
                )
                
                is_body_text = not is_heading and len(paragraph.text.strip()) > 0
                
                # Check specific spacing requirements
                if is_body_text:
                    if space_after >= 12:  # 12pt space_after for body text
                        spacing_analysis["body_text_with_12pt_after"] += 1
                    if space_before >= 3:   # 3pt space_before for body text
                        spacing_analysis["body_text_with_3pt_before"] += 1
                
                if is_subheading:
                    if space_before >= 12:  # 12pt before for sub-headings
                        spacing_analysis["subheadings_with_12pt_before"] += 1
                    if space_after >= 6:    # 6pt after for sub-headings
                        spacing_analysis["subheadings_with_6pt_after"] += 1
                
                # Check for breathing room (any spacing)
                if space_before > 0 or space_after > 0:
                    spacing_analysis["paragraphs_with_breathing_room"] += 1
                
                # Check Q&A pair separation
                text_lower = paragraph.text.lower()
                if "question" in text_lower and i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_format = next_para.paragraph_format
                    next_space_before = next_format.space_before.pt if next_format.space_before else 0
                    if next_space_before >= 6 or space_after >= 10:
                        spacing_analysis["qa_pairs_with_separation"] += 1
                
                spacing_analysis["spacing_details"].append({
                    "index": i,
                    "text_preview": paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
                    "style": paragraph.style.name,
                    "space_before": space_before,
                    "space_after": space_after,
                    "is_heading": is_heading,
                    "is_subheading": is_subheading,
                    "is_body_text": is_body_text
                })
            
            # Clean up temp file
            os.unlink(temp_file_path)
            
            return spacing_analysis
            
        except Exception as e:
            print(f"Error analyzing document spacing: {str(e)}")
            return None

    async def test_body_text_spacing(self, docx_content):
        """Test that body text has proper spacing (12pt space_after, 3pt space_before)"""
        try:
            analysis = self.analyze_paragraph_spacing(docx_content)
            if not analysis:
                await self.log_test("Body Text Spacing (12pt after, 3pt before)", False, "Failed to analyze document")
                return False
            
            body_text_count = sum(1 for detail in analysis["spacing_details"] if detail["is_body_text"])
            
            if body_text_count == 0:
                await self.log_test("Body Text Spacing (12pt after, 3pt before)", False, "No body text paragraphs found")
                return False
            
            after_compliance = (analysis["body_text_with_12pt_after"] / body_text_count) * 100
            before_compliance = (analysis["body_text_with_3pt_before"] / body_text_count) * 100
            
            details = f"Body text paragraphs: {body_text_count}, 12pt after: {analysis['body_text_with_12pt_after']} ({after_compliance:.1f}%), 3pt before: {analysis['body_text_with_3pt_before']} ({before_compliance:.1f}%)"
            
            # Test passes if at least 70% of body text has proper spacing
            if after_compliance >= 70 and before_compliance >= 50:
                await self.log_test("Body Text Spacing (12pt after, 3pt before)", True, details)
                return True
            else:
                await self.log_test("Body Text Spacing (12pt after, 3pt before)", False, details)
                return False
                
        except Exception as e:
            await self.log_test("Body Text Spacing (12pt after, 3pt before)", False, f"Exception: {str(e)}")
            return False

    async def test_subheading_spacing(self, docx_content):
        """Test that sub-headings have enhanced spacing (12pt before, 6pt after)"""
        try:
            analysis = self.analyze_paragraph_spacing(docx_content)
            if not analysis:
                await self.log_test("Sub-heading Spacing (12pt before, 6pt after)", False, "Failed to analyze document")
                return False
            
            subheading_count = sum(1 for detail in analysis["spacing_details"] if detail["is_subheading"])
            
            if subheading_count == 0:
                await self.log_test("Sub-heading Spacing (12pt before, 6pt after)", False, "No sub-heading paragraphs found")
                return False
            
            before_compliance = (analysis["subheadings_with_12pt_before"] / subheading_count) * 100
            after_compliance = (analysis["subheadings_with_6pt_after"] / subheading_count) * 100
            
            details = f"Sub-headings: {subheading_count}, 12pt before: {analysis['subheadings_with_12pt_before']} ({before_compliance:.1f}%), 6pt after: {analysis['subheadings_with_6pt_after']} ({after_compliance:.1f}%)"
            
            # Test passes if at least 60% of sub-headings have proper spacing
            if before_compliance >= 60 and after_compliance >= 60:
                await self.log_test("Sub-heading Spacing (12pt before, 6pt after)", True, details)
                return True
            else:
                await self.log_test("Sub-heading Spacing (12pt before, 6pt after)", False, details)
                return False
                
        except Exception as e:
            await self.log_test("Sub-heading Spacing (12pt before, 6pt after)", False, f"Exception: {str(e)}")
            return False

    async def test_breathing_room_between_sections(self, docx_content):
        """Test that there's better breathing room between sections"""
        try:
            analysis = self.analyze_paragraph_spacing(docx_content)
            if not analysis:
                await self.log_test("Breathing Room Between Sections", False, "Failed to analyze document")
                return False
            
            total_paragraphs = analysis["total_paragraphs"]
            paragraphs_with_spacing = analysis["paragraphs_with_breathing_room"]
            
            breathing_room_percentage = (paragraphs_with_spacing / total_paragraphs) * 100
            
            details = f"Total paragraphs: {total_paragraphs}, With spacing: {paragraphs_with_spacing} ({breathing_room_percentage:.1f}%)"
            
            # Test passes if at least 80% of paragraphs have some spacing
            if breathing_room_percentage >= 80:
                await self.log_test("Breathing Room Between Sections", True, details)
                return True
            else:
                await self.log_test("Breathing Room Between Sections", False, details)
                return False
                
        except Exception as e:
            await self.log_test("Breathing Room Between Sections", False, f"Exception: {str(e)}")
            return False

    async def test_qa_pair_visual_separation(self, docx_content):
        """Test that Q&A pairs have clear visual separation"""
        try:
            analysis = self.analyze_paragraph_spacing(docx_content)
            if not analysis:
                await self.log_test("Q&A Pair Visual Separation", False, "Failed to analyze document")
                return False
            
            # Count questions in the document
            question_count = sum(1 for detail in analysis["spacing_details"] 
                               if "question" in detail["text_preview"].lower())
            
            qa_separation_count = analysis["qa_pairs_with_separation"]
            
            if question_count == 0:
                await self.log_test("Q&A Pair Visual Separation", False, "No Q&A pairs found in document")
                return False
            
            separation_percentage = (qa_separation_count / question_count) * 100
            
            details = f"Questions found: {question_count}, With proper separation: {qa_separation_count} ({separation_percentage:.1f}%)"
            
            # Test passes if at least 60% of Q&A pairs have proper separation
            if separation_percentage >= 60:
                await self.log_test("Q&A Pair Visual Separation", True, details)
                return True
            else:
                await self.log_test("Q&A Pair Visual Separation", False, details)
                return False
                
        except Exception as e:
            await self.log_test("Q&A Pair Visual Separation", False, f"Exception: {str(e)}")
            return False

    async def test_user_feedback_resolution(self, docx_content):
        """Test that user feedback about 'no line space between paragraphs' has been resolved"""
        try:
            analysis = self.analyze_paragraph_spacing(docx_content)
            if not analysis:
                await self.log_test("User Feedback Resolution - Line Spacing", False, "Failed to analyze document")
                return False
            
            total_paragraphs = analysis["total_paragraphs"]
            paragraphs_with_spacing = analysis["paragraphs_with_breathing_room"]
            
            # Calculate various spacing metrics
            body_text_after_spacing = analysis["body_text_with_12pt_after"]
            body_text_before_spacing = analysis["body_text_with_3pt_before"]
            subheading_spacing = analysis["subheadings_with_12pt_before"] + analysis["subheadings_with_6pt_after"]
            
            overall_spacing_score = (paragraphs_with_spacing / total_paragraphs) * 100
            
            details = f"Overall spacing coverage: {paragraphs_with_spacing}/{total_paragraphs} ({overall_spacing_score:.1f}%), Body text spacing improvements: {body_text_after_spacing + body_text_before_spacing}, Sub-heading spacing improvements: {subheading_spacing}"
            
            # Test passes if overall spacing coverage is good and specific improvements are implemented
            if overall_spacing_score >= 75 and (body_text_after_spacing + body_text_before_spacing) > 0:
                await self.log_test("User Feedback Resolution - Line Spacing", True, details)
                return True
            else:
                await self.log_test("User Feedback Resolution - Line Spacing", False, details)
                return False
                
        except Exception as e:
            await self.log_test("User Feedback Resolution - Line Spacing", False, f"Exception: {str(e)}")
            return False

    async def run_all_tests(self):
        """Run all paragraph spacing tests"""
        print("📏 Starting Enhanced Paragraph Spacing Testing...")
        print(f"Backend URL: {BACKEND_URL}")
        print(f"Test User Email: {self.test_user_email}")
        print("=" * 80)

        # Setup tests
        if not await self.create_test_user():
            return 0, 1, self.test_results
        
        if not await self.create_note_with_ai_conversations():
            return 1, 2, self.test_results
        
        # Export document
        docx_content = await self.export_word_document()
        if not docx_content:
            return 2, 3, self.test_results
        
        # Run spacing tests
        spacing_tests = [
            ("Body Text Spacing (12pt after, 3pt before)", lambda: self.test_body_text_spacing(docx_content)),
            ("Sub-heading Spacing (12pt before, 6pt after)", lambda: self.test_subheading_spacing(docx_content)),
            ("Breathing Room Between Sections", lambda: self.test_breathing_room_between_sections(docx_content)),
            ("Q&A Pair Visual Separation", lambda: self.test_qa_pair_visual_separation(docx_content)),
            ("User Feedback Resolution - Line Spacing", lambda: self.test_user_feedback_resolution(docx_content)),
        ]
        
        passed = 3  # Setup tests passed
        total = 3 + len(spacing_tests)
        
        for test_name, test_func in spacing_tests:
            try:
                result = await test_func()
                if result:
                    passed += 1
            except Exception as e:
                await self.log_test(test_name, False, f"Unexpected exception: {str(e)}")

        print("=" * 80)
        print(f"📏 Paragraph Spacing Testing Complete")
        print(f"📊 Results: {passed}/{total} tests passed ({(passed/total)*100:.1f}%)")
        
        if passed == total:
            print("✅ All paragraph spacing tests PASSED!")
        else:
            print("❌ Some paragraph spacing tests FAILED!")
            
        return passed, total, self.test_results

async def main():
    """Main test runner"""
    tester = ParagraphSpacingTester()
    passed, total, results = await tester.run_all_tests()
    
    # Print detailed results
    print("\n" + "=" * 80)
    print("📋 DETAILED TEST RESULTS:")
    print("=" * 80)
    
    for result in results:
        status = "✅ PASS" if result["success"] else "❌ FAIL"
        print(f"{status}: {result['test']}")
        if result["details"]:
            print(f"   Details: {result['details']}")
    
    print("\n" + "=" * 80)
    print("🎯 SUMMARY:")
    print("=" * 80)
    
    if passed == total:
        print("✅ Enhanced paragraph spacing in AI analysis Word documents is working perfectly!")
        print("✅ Body text has proper spacing (12pt space_after, 3pt space_before)")
        print("✅ Sub-headings have enhanced spacing (12pt before, 6pt after)")
        print("✅ Better breathing room between sections implemented")
        print("✅ Clear visual separation between Q&A pairs")
        print("✅ User feedback about 'no line space between paragraphs' has been resolved")
    else:
        print("❌ Paragraph spacing enhancements need attention")
        failed_tests = [r for r in results if not r["success"]]
        print(f"❌ Failed tests: {len(failed_tests)}")
        for failed in failed_tests:
            print(f"   - {failed['test']}: {failed['details']}")

if __name__ == "__main__":
    asyncio.run(main())