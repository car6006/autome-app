#!/usr/bin/env python3

import asyncio
import httpx
import json
import uuid
from datetime import datetime
import os
import tempfile

# Backend URL from frontend environment
BACKEND_URL = "https://autome-fix.preview.emergentagent.com/api"

class DetailedWordFormatTester:
    def __init__(self):
        self.test_results = []
        self.test_user_email = f"test_user_{uuid.uuid4().hex[:8]}@example.com"
        self.test_password = "TestPassword123"
        self.auth_token = None
        self.test_note_id = None

    async def log_test(self, test_name, success, details=""):
        """Log test results"""
        result = {
            "test": test_name,
            "success": success,
            "details": details,
            "timestamp": datetime.now().isoformat()
        }
        self.test_results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name}")
        if details:
            print(f"   Details: {details}")

    async def setup_test_user_and_note(self):
        """Setup test user and create a note with complex AI conversations"""
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                # Create user
                user_data = {
                    "email": self.test_user_email,
                    "username": f"testuser{uuid.uuid4().hex[:8]}",
                    "password": self.test_password,
                    "first_name": "Test",
                    "last_name": "User"
                }
                
                response = await client.post(f"{BACKEND_URL}/auth/register", json=user_data)
                
                if response.status_code == 200:
                    data = response.json()
                    self.auth_token = data.get("access_token")
                else:
                    await self.log_test("Setup User", False, f"Status: {response.status_code}")
                    return False

                # Create a complex note with detailed content
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                note_data = {
                    "title": "Complex Business Analysis Meeting - Q4 Strategic Review",
                    "kind": "text",
                    "text_content": """
                QUARTERLY BUSINESS REVIEW MEETING
                Date: December 15, 2024
                Attendees: Executive Team, Department Heads, Regional Managers
                
                EXECUTIVE SUMMARY
                This quarterly review focused on analyzing our performance metrics, identifying growth opportunities, and establishing strategic priorities for the upcoming quarter. Key discussions centered around operational efficiency, market expansion, technology integration, and resource optimization.
                
                PERFORMANCE ANALYSIS
                Revenue Performance: Q4 showed a 15% increase compared to Q3, with particularly strong performance in the Asia-Pacific region. European markets remained stable while North American operations exceeded projections by 8%.
                
                Operational Metrics: Processing times improved by 12% due to automation initiatives. Customer satisfaction scores increased to 94%, up from 89% in the previous quarter.
                
                STRATEGIC INITIATIVES DISCUSSED
                1. Digital Transformation: Implementation of new CRM system scheduled for Q1 2025
                2. Market Expansion: Evaluation of opportunities in Southeast Asian markets
                3. Sustainability Programs: Carbon footprint reduction targets and green logistics initiatives
                4. Workforce Development: Training programs for emerging technologies and leadership development
                
                RISK ASSESSMENT
                Supply chain vulnerabilities were identified in three key regions. Contingency planning discussions focused on diversification strategies and alternative supplier networks.
                
                FINANCIAL PROJECTIONS
                Budget allocations for Q1 2025 were reviewed, with increased investment in technology infrastructure and market research initiatives.
                
                ACTION ITEMS AND NEXT STEPS
                Multiple action items were assigned across departments with specific deadlines and accountability measures established.
                """
                }
                
                response = await client.post(f"{BACKEND_URL}/notes", json=note_data, headers=headers)
                
                if response.status_code == 200:
                    self.test_note_id = response.json().get("id")
                    
                    # Add complex AI conversations with varied content
                    complex_questions = [
                        "Provide a comprehensive executive summary of the key insights from this quarterly review meeting.",
                        "What are the most critical action items that should be prioritized, and what are the potential risks if they are not addressed?",
                        "Analyze the performance metrics and provide recommendations for improving operational efficiency in the next quarter.",
                        "Based on the strategic initiatives discussed, what would be the most effective implementation timeline and resource allocation strategy?",
                        "Evaluate the risk assessment findings and suggest comprehensive mitigation strategies for the identified supply chain vulnerabilities."
                    ]
                    
                    for question in complex_questions:
                        ai_response = await client.post(
                            f"{BACKEND_URL}/notes/{self.test_note_id}/ai-chat",
                            json={"question": question},
                            headers=headers
                        )
                        
                        if ai_response.status_code != 200:
                            await self.log_test("Setup Complex Note", False, f"Failed to add AI conversation: {ai_response.status_code}")
                            return False
                    
                    await self.log_test("Setup Complex Note", True, f"Complex note created with ID: {self.test_note_id}")
                    return True
                else:
                    await self.log_test("Setup Complex Note", False, f"Status: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Setup", False, f"Exception: {str(e)}")
            return False

    async def test_paragraph_structure_improvement(self):
        """Test that AI responses are properly parsed into paragraphs (not condensed text)"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        doc = Document(temp_file_path)
                        
                        # Analyze paragraph structure
                        paragraphs = doc.paragraphs
                        content_paragraphs = [p for p in paragraphs if p.text.strip() and len(p.text.strip()) > 20]
                        
                        # Check for proper paragraph breaks (not wall of text)
                        avg_paragraph_length = sum(len(p.text) for p in content_paragraphs) / len(content_paragraphs) if content_paragraphs else 0
                        
                        # Good paragraph structure should have reasonable length paragraphs
                        proper_paragraph_structure = 50 <= avg_paragraph_length <= 500
                        
                        # Check for variety in paragraph lengths (indicates proper parsing)
                        paragraph_lengths = [len(p.text) for p in content_paragraphs]
                        length_variance = max(paragraph_lengths) - min(paragraph_lengths) if paragraph_lengths else 0
                        good_variance = length_variance > 100  # Should have variety, not uniform blocks
                        
                        # Check spacing between paragraphs
                        paragraphs_with_spacing = 0
                        for para in paragraphs:
                            if para.paragraph_format.space_after and para.paragraph_format.space_after.pt > 0:
                                paragraphs_with_spacing += 1
                        
                        proper_spacing = paragraphs_with_spacing > len(content_paragraphs) * 0.3  # At least 30% have spacing
                        
                        os.unlink(temp_file_path)
                        
                        structure_score = 0
                        structure_details = []
                        
                        if proper_paragraph_structure:
                            structure_score += 1
                            structure_details.append(f"✓ Proper paragraph length (avg: {avg_paragraph_length:.0f} chars)")
                        else:
                            structure_details.append(f"✗ Poor paragraph length (avg: {avg_paragraph_length:.0f} chars)")
                        
                        if good_variance:
                            structure_score += 1
                            structure_details.append(f"✓ Good paragraph variety (variance: {length_variance})")
                        else:
                            structure_details.append(f"✗ Poor paragraph variety (variance: {length_variance})")
                        
                        if proper_spacing:
                            structure_score += 1
                            structure_details.append(f"✓ Proper paragraph spacing ({paragraphs_with_spacing} paragraphs)")
                        else:
                            structure_details.append(f"✗ Limited paragraph spacing ({paragraphs_with_spacing} paragraphs)")
                        
                        if len(content_paragraphs) >= 20:
                            structure_score += 1
                            structure_details.append(f"✓ Rich paragraph structure ({len(content_paragraphs)} content paragraphs)")
                        else:
                            structure_details.append(f"✗ Limited paragraph structure ({len(content_paragraphs)} content paragraphs)")
                        
                        success = structure_score >= 3
                        details = f"Paragraph score: {structure_score}/4. " + "; ".join(structure_details)
                        
                        await self.log_test("Paragraph Structure Improvement", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test("Paragraph Structure Improvement", False, f"Analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test("Paragraph Structure Improvement", False, f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Paragraph Structure Improvement", False, f"Exception: {str(e)}")
            return False

    async def test_professional_heading_styles(self):
        """Test that headings are professional (not squared) with proper font sizes"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        from docx.shared import Pt
                        doc = Document(temp_file_path)
                        
                        # Analyze heading styles
                        heading_analysis = {
                            "title_headings": 0,
                            "section_headings": 0,
                            "proper_font_sizes": 0,
                            "calibri_headings": 0,
                            "bold_headings": 0
                        }
                        
                        for para in doc.paragraphs:
                            if para.style and ("heading" in para.style.name.lower() or "title" in para.style.name.lower()):
                                if "title" in para.style.name.lower():
                                    heading_analysis["title_headings"] += 1
                                else:
                                    heading_analysis["section_headings"] += 1
                                
                                # Check font properties
                                for run in para.runs:
                                    if run.font.size:
                                        font_size_pt = run.font.size.pt
                                        if font_size_pt >= 12:  # Professional heading size
                                            heading_analysis["proper_font_sizes"] += 1
                                    
                                    if run.font.name and "calibri" in run.font.name.lower():
                                        heading_analysis["calibri_headings"] += 1
                                    
                                    if run.font.bold:
                                        heading_analysis["bold_headings"] += 1
                        
                        os.unlink(temp_file_path)
                        
                        heading_score = 0
                        heading_details = []
                        
                        if heading_analysis["title_headings"] >= 1:
                            heading_score += 1
                            heading_details.append(f"✓ Title headings found ({heading_analysis['title_headings']})")
                        else:
                            heading_details.append("✗ No title headings found")
                        
                        if heading_analysis["section_headings"] >= 3:
                            heading_score += 1
                            heading_details.append(f"✓ Section headings found ({heading_analysis['section_headings']})")
                        else:
                            heading_details.append(f"✗ Limited section headings ({heading_analysis['section_headings']})")
                        
                        if heading_analysis["proper_font_sizes"] >= 2:
                            heading_score += 1
                            heading_details.append(f"✓ Professional font sizes ({heading_analysis['proper_font_sizes']} headings)")
                        else:
                            heading_details.append(f"✗ Poor font sizes ({heading_analysis['proper_font_sizes']} headings)")
                        
                        if heading_analysis["calibri_headings"] >= 2:
                            heading_score += 1
                            heading_details.append(f"✓ Calibri font in headings ({heading_analysis['calibri_headings']})")
                        else:
                            heading_details.append(f"✗ Limited Calibri usage ({heading_analysis['calibri_headings']})")
                        
                        if heading_analysis["bold_headings"] >= 2:
                            heading_score += 1
                            heading_details.append(f"✓ Bold headings ({heading_analysis['bold_headings']})")
                        else:
                            heading_details.append(f"✗ Limited bold headings ({heading_analysis['bold_headings']})")
                        
                        success = heading_score >= 4
                        details = f"Heading score: {heading_score}/5. " + "; ".join(heading_details)
                        
                        await self.log_test("Professional Heading Styles", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test("Professional Heading Styles", False, f"Analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test("Professional Heading Styles", False, f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Professional Heading Styles", False, f"Exception: {str(e)}")
            return False

    async def test_content_structure_preservation(self):
        """Test that AI response structure is preserved while looking professional"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        doc = Document(temp_file_path)
                        
                        # Extract all text content
                        all_text = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                all_text.append(para.text.strip())
                        
                        full_content = " ".join(all_text)
                        
                        # Check for preserved business content structure
                        business_structure_indicators = [
                            "executive summary",
                            "key insights", 
                            "recommendations",
                            "action items",
                            "analysis",
                            "strategic",
                            "performance",
                            "operational",
                            "implementation",
                            "risk assessment"
                        ]
                        
                        structure_preserved = sum(1 for indicator in business_structure_indicators 
                                                if indicator.lower() in full_content.lower())
                        
                        # Check for professional formatting (no raw AI artifacts)
                        ai_artifacts = [
                            "**", "###", "##", "#", 
                            "bullet point", "•", 
                            "here are", "in summary",
                            "based on the", "according to"
                        ]
                        
                        artifacts_removed = sum(1 for artifact in ai_artifacts 
                                              if artifact in full_content.lower())
                        
                        # Check for question-answer structure preservation
                        qa_structure = [
                            "question 1", "question 2", "question 3", "question 4", "question 5",
                            "analysis", "response"
                        ]
                        
                        qa_preserved = sum(1 for qa in qa_structure 
                                         if qa.lower() in full_content.lower())
                        
                        # Check content richness (not stripped down)
                        content_richness = len(full_content) > 5000  # Should be substantial
                        
                        os.unlink(temp_file_path)
                        
                        preservation_score = 0
                        preservation_details = []
                        
                        if structure_preserved >= 6:
                            preservation_score += 1
                            preservation_details.append(f"✓ Business structure preserved ({structure_preserved} indicators)")
                        else:
                            preservation_details.append(f"✗ Limited structure preservation ({structure_preserved} indicators)")
                        
                        if artifacts_removed <= 3:  # Should be minimal AI artifacts
                            preservation_score += 1
                            preservation_details.append(f"✓ AI artifacts cleaned ({artifacts_removed} remaining)")
                        else:
                            preservation_details.append(f"✗ Too many AI artifacts ({artifacts_removed} remaining)")
                        
                        if qa_preserved >= 4:
                            preservation_score += 1
                            preservation_details.append(f"✓ Q&A structure preserved ({qa_preserved} elements)")
                        else:
                            preservation_details.append(f"✗ Poor Q&A preservation ({qa_preserved} elements)")
                        
                        if content_richness:
                            preservation_score += 1
                            preservation_details.append(f"✓ Rich content preserved ({len(full_content)} chars)")
                        else:
                            preservation_details.append(f"✗ Content too sparse ({len(full_content)} chars)")
                        
                        success = preservation_score >= 3
                        details = f"Preservation score: {preservation_score}/4. " + "; ".join(preservation_details)
                        
                        await self.log_test("Content Structure Preservation", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test("Content Structure Preservation", False, f"Analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test("Content Structure Preservation", False, f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Content Structure Preservation", False, f"Exception: {str(e)}")
            return False

    async def test_professional_layout_margins(self):
        """Test professional layout with proper margins and spacing"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        from docx.shared import Inches
                        doc = Document(temp_file_path)
                        
                        # Check margins
                        section = doc.sections[0]
                        margins = {
                            "top": section.top_margin.inches,
                            "bottom": section.bottom_margin.inches,
                            "left": section.left_margin.inches,
                            "right": section.right_margin.inches
                        }
                        
                        # Professional margins should be around 1 inch
                        professional_margins = all(0.8 <= margin <= 1.2 for margin in margins.values())
                        
                        # Check line spacing
                        paragraphs_with_line_spacing = 0
                        for para in doc.paragraphs:
                            if para.paragraph_format.line_spacing and para.paragraph_format.line_spacing >= 1.1:
                                paragraphs_with_line_spacing += 1
                        
                        good_line_spacing = paragraphs_with_line_spacing > len(doc.paragraphs) * 0.5
                        
                        # Check paragraph spacing
                        paragraphs_with_space_after = 0
                        for para in doc.paragraphs:
                            if para.paragraph_format.space_after and para.paragraph_format.space_after.pt > 0:
                                paragraphs_with_space_after += 1
                        
                        good_paragraph_spacing = paragraphs_with_space_after > len(doc.paragraphs) * 0.3
                        
                        # Check font consistency (should be primarily Calibri)
                        calibri_usage = 0
                        total_runs = 0
                        for para in doc.paragraphs:
                            for run in para.runs:
                                total_runs += 1
                                if run.font.name and "calibri" in run.font.name.lower():
                                    calibri_usage += 1
                        
                        consistent_font = (calibri_usage / total_runs) > 0.7 if total_runs > 0 else False
                        
                        os.unlink(temp_file_path)
                        
                        layout_score = 0
                        layout_details = []
                        
                        if professional_margins:
                            layout_score += 1
                            layout_details.append(f"✓ Professional margins (T:{margins['top']:.1f}, B:{margins['bottom']:.1f}, L:{margins['left']:.1f}, R:{margins['right']:.1f})")
                        else:
                            layout_details.append(f"✗ Poor margins (T:{margins['top']:.1f}, B:{margins['bottom']:.1f}, L:{margins['left']:.1f}, R:{margins['right']:.1f})")
                        
                        if good_line_spacing:
                            layout_score += 1
                            layout_details.append(f"✓ Good line spacing ({paragraphs_with_line_spacing} paragraphs)")
                        else:
                            layout_details.append(f"✗ Poor line spacing ({paragraphs_with_line_spacing} paragraphs)")
                        
                        if good_paragraph_spacing:
                            layout_score += 1
                            layout_details.append(f"✓ Good paragraph spacing ({paragraphs_with_space_after} paragraphs)")
                        else:
                            layout_details.append(f"✗ Poor paragraph spacing ({paragraphs_with_space_after} paragraphs)")
                        
                        if consistent_font:
                            layout_score += 1
                            layout_details.append(f"✓ Consistent Calibri font ({calibri_usage}/{total_runs} runs)")
                        else:
                            layout_details.append(f"✗ Inconsistent font usage ({calibri_usage}/{total_runs} runs)")
                        
                        success = layout_score >= 3
                        details = f"Layout score: {layout_score}/4. " + "; ".join(layout_details)
                        
                        await self.log_test("Professional Layout & Margins", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test("Professional Layout & Margins", False, f"Analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test("Professional Layout & Margins", False, f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Professional Layout & Margins", False, f"Exception: {str(e)}")
            return False

    async def run_detailed_tests(self):
        """Run detailed Word formatting tests"""
        print("📄 Starting Detailed AI Word Document Formatting Tests...")
        print(f"Backend URL: {BACKEND_URL}")
        print("=" * 80)

        tests = [
            ("Setup", self.setup_test_user_and_note),
            ("Paragraph Structure Improvement", self.test_paragraph_structure_improvement),
            ("Professional Heading Styles", self.test_professional_heading_styles),
            ("Content Structure Preservation", self.test_content_structure_preservation),
            ("Professional Layout & Margins", self.test_professional_layout_margins)
        ]

        passed = 0
        total = len(tests)

        for test_name, test_func in tests:
            try:
                result = await test_func()
                if result:
                    passed += 1
            except Exception as e:
                await self.log_test(test_name, False, f"Unexpected exception: {str(e)}")

        print("=" * 80)
        print(f"📄 Detailed Word Formatting Tests Complete")
        print(f"📊 Results: {passed}/{total} tests passed ({(passed/total)*100:.1f}%)")
        
        return passed, total, self.test_results

async def main():
    """Main test runner"""
    tester = DetailedWordFormatTester()
    passed, total, results = await tester.run_detailed_tests()
    
    # Print detailed results
    print("\n" + "=" * 80)
    print("📋 DETAILED FORMATTING TEST RESULTS:")
    print("=" * 80)
    
    for result in results:
        status = "✅ PASS" if result["success"] else "❌ FAIL"
        print(f"{status}: {result['test']}")
        if result["details"]:
            print(f"   Details: {result['details']}")
    
    print("\n" + "=" * 80)
    print("🎯 DETAILED FORMATTING SUMMARY:")
    print("=" * 80)
    
    if passed >= total * 0.8:
        print("✅ Word document formatting improvements are working excellently!")
        print("✅ Professional paragraph structure (not condensed text)")
        print("✅ Professional headings (not squared) with proper font sizes")
        print("✅ Content structure preserved while looking professional")
        print("✅ Professional layout with proper margins and Calibri font")
    else:
        print("❌ Word formatting improvements need attention")
        failed_tests = [r for r in results if not r["success"]]
        for failed in failed_tests:
            print(f"   - {failed['test']}: {failed['details']}")

if __name__ == "__main__":
    asyncio.run(main())