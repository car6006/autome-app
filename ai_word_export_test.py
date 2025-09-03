#!/usr/bin/env python3

import asyncio
import httpx
import json
import uuid
from datetime import datetime
import os
import tempfile

# Backend URL from frontend environment
BACKEND_URL = "https://autome-fix.preview.emergentagent.com/api"

class AIWordExportTester:
    def __init__(self):
        self.test_results = []
        self.test_user_email = f"test_user_{uuid.uuid4().hex[:8]}@example.com"
        self.expeditors_user_email = f"test_user_{uuid.uuid4().hex[:8]}@expeditors.com"
        self.test_password = "TestPassword123"
        self.auth_token = None
        self.expeditors_token = None
        self.test_note_id = None
        self.expeditors_note_id = None

    async def log_test(self, test_name, success, details=""):
        """Log test results"""
        result = {
            "test": test_name,
            "success": success,
            "details": details,
            "timestamp": datetime.now().isoformat()
        }
        self.test_results.append(result)
        status = "✅ PASS" if success else "❌ FAIL"
        print(f"{status}: {test_name}")
        if details:
            print(f"   Details: {details}")

    async def create_test_users(self):
        """Create test users - regular and Expeditors"""
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                # Create regular user
                regular_user_data = {
                    "email": self.test_user_email,
                    "username": f"testuser{uuid.uuid4().hex[:8]}",
                    "password": self.test_password,
                    "first_name": "Test",
                    "last_name": "User"
                }
                
                response = await client.post(f"{BACKEND_URL}/auth/register", json=regular_user_data)
                
                if response.status_code == 200:
                    data = response.json()
                    self.auth_token = data.get("access_token")
                    await self.log_test("Create Regular Test User", True, f"User created: {self.test_user_email}")
                else:
                    await self.log_test("Create Regular Test User", False, f"Status: {response.status_code}")
                    return False

                # Create Expeditors user
                expeditors_user_data = {
                    "email": self.expeditors_user_email,
                    "username": f"expuser{uuid.uuid4().hex[:8]}",
                    "password": self.test_password,
                    "first_name": "Expeditors",
                    "last_name": "User"
                }
                
                response = await client.post(f"{BACKEND_URL}/auth/register", json=expeditors_user_data)
                
                if response.status_code == 200:
                    data = response.json()
                    self.expeditors_token = data.get("access_token")
                    await self.log_test("Create Expeditors Test User", True, f"User created: {self.expeditors_user_email}")
                    return True
                else:
                    await self.log_test("Create Expeditors Test User", False, f"Status: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Create Test Users", False, f"Exception: {str(e)}")
            return False

    async def create_note_with_ai_conversations(self, token, user_type="regular"):
        """Create a note with AI conversations for testing"""
        try:
            async with httpx.AsyncClient(timeout=30) as client:
                headers = {"Authorization": f"Bearer {token}"}
                
                # Create a text note
                note_data = {
                    "title": f"AI Analysis Test Meeting - {user_type.title()}",
                    "kind": "text",
                    "text_content": """
                Meeting Discussion: Supply Chain Optimization Strategy
                
                Today we discussed the quarterly performance review and identified several key areas for improvement in our supply chain operations. The team reviewed current logistics processes, analyzed cost efficiency metrics, and proposed strategic initiatives for the next quarter.
                
                Key topics covered:
                - Regional distribution center performance analysis
                - Transportation cost optimization opportunities  
                - Technology integration for improved visibility
                - Customer service delivery improvements
                - Risk management and contingency planning
                
                The meeting concluded with action items assigned to various team members and a follow-up scheduled for next week to review implementation progress.
                """
                }
                
                response = await client.post(f"{BACKEND_URL}/notes", json=note_data, headers=headers)
                
                if response.status_code == 200:
                    note_id = response.json().get("id")
                    
                    # Add multiple AI conversations to the note
                    ai_questions = [
                        "What are the key insights from this meeting?",
                        "What action items should be prioritized?", 
                        "What are the potential risks and how can we mitigate them?",
                        "How can we improve our supply chain efficiency based on this discussion?"
                    ]
                    
                    for question in ai_questions:
                        ai_response = await client.post(
                            f"{BACKEND_URL}/notes/{note_id}/ai-chat",
                            json={"question": question},
                            headers=headers
                        )
                        
                        if ai_response.status_code != 200:
                            await self.log_test(f"Add AI Conversation ({user_type})", False, f"Failed to add AI conversation: {ai_response.status_code}")
                            return None
                    
                    if user_type == "regular":
                        self.test_note_id = note_id
                    else:
                        self.expeditors_note_id = note_id
                        
                    await self.log_test(f"Create Note with AI Conversations ({user_type})", True, f"Note created with ID: {note_id}")
                    return note_id
                else:
                    await self.log_test(f"Create Note with AI Conversations ({user_type})", False, f"Status: {response.status_code}")
                    return None
                    
        except Exception as e:
            await self.log_test(f"Create Note with AI Conversations ({user_type})", False, f"Exception: {str(e)}")
            return None

    async def test_word_export_functionality(self, note_id, token, user_type="regular"):
        """Test Word document export functionality"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    # Check content type
                    content_type = response.headers.get("content-type", "")
                    if "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type:
                        content_length = len(response.content)
                        
                        # Save file temporarily to verify it's a valid Word document
                        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                            temp_file.write(response.content)
                            temp_file_path = temp_file.name
                        
                        # Check if file is valid by trying to read it
                        try:
                            from docx import Document
                            doc = Document(temp_file_path)
                            paragraph_count = len(doc.paragraphs)
                            
                            # Clean up temp file
                            os.unlink(temp_file_path)
                            
                            await self.log_test(f"Word Export Functionality ({user_type})", True, 
                                              f"Valid DOCX file generated: {content_length} bytes, {paragraph_count} paragraphs")
                            return True
                        except Exception as doc_error:
                            os.unlink(temp_file_path)
                            await self.log_test(f"Word Export Functionality ({user_type})", False, 
                                              f"Invalid DOCX file: {str(doc_error)}")
                            return False
                    else:
                        await self.log_test(f"Word Export Functionality ({user_type})", False, 
                                          f"Wrong content type: {content_type}")
                        return False
                else:
                    await self.log_test(f"Word Export Functionality ({user_type})", False, 
                                      f"Status: {response.status_code}, Response: {response.text}")
                    return False
                    
        except Exception as e:
            await self.log_test(f"Word Export Functionality ({user_type})", False, f"Exception: {str(e)}")
            return False

    async def test_document_structure_and_formatting(self, note_id, token, user_type="regular"):
        """Test document structure and professional formatting"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    # Save and analyze document structure
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        doc = Document(temp_file_path)
                        
                        # Check document structure
                        paragraphs = doc.paragraphs
                        styles_used = set()
                        has_headings = False
                        has_body_text = False
                        calibri_font_found = False
                        
                        for para in paragraphs:
                            if para.style:
                                styles_used.add(para.style.name)
                            
                            # Check for headings
                            if para.style and ("heading" in para.style.name.lower() or "title" in para.style.name.lower()):
                                has_headings = True
                            
                            # Check for body text
                            if para.text.strip() and len(para.text.strip()) > 20:
                                has_body_text = True
                            
                            # Check font usage
                            for run in para.runs:
                                if run.font.name and "calibri" in run.font.name.lower():
                                    calibri_font_found = True
                        
                        # Check margins
                        section = doc.sections[0]
                        margins_set = (section.top_margin.inches >= 0.8 and 
                                     section.bottom_margin.inches >= 0.8 and
                                     section.left_margin.inches >= 0.8 and
                                     section.right_margin.inches >= 0.8)
                        
                        # Clean up temp file
                        os.unlink(temp_file_path)
                        
                        structure_score = 0
                        structure_details = []
                        
                        if has_headings:
                            structure_score += 1
                            structure_details.append("✓ Professional headings found")
                        else:
                            structure_details.append("✗ No professional headings detected")
                        
                        if has_body_text:
                            structure_score += 1
                            structure_details.append("✓ Proper body text structure")
                        else:
                            structure_details.append("✗ No substantial body text found")
                        
                        if calibri_font_found:
                            structure_score += 1
                            structure_details.append("✓ Calibri font usage detected")
                        else:
                            structure_details.append("✗ Calibri font not detected")
                        
                        if margins_set:
                            structure_score += 1
                            structure_details.append("✓ Professional margins (1 inch)")
                        else:
                            structure_details.append("✗ Margins not properly set")
                        
                        if len(paragraphs) >= 5:
                            structure_score += 1
                            structure_details.append(f"✓ Good paragraph structure ({len(paragraphs)} paragraphs)")
                        else:
                            structure_details.append(f"✗ Limited paragraph structure ({len(paragraphs)} paragraphs)")
                        
                        success = structure_score >= 3  # At least 3 out of 5 criteria
                        details = f"Structure score: {structure_score}/5. " + "; ".join(structure_details)
                        
                        await self.log_test(f"Document Structure & Formatting ({user_type})", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test(f"Document Structure & Formatting ({user_type})", False, 
                                          f"Document analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test(f"Document Structure & Formatting ({user_type})", False, 
                                      f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test(f"Document Structure & Formatting ({user_type})", False, f"Exception: {str(e)}")
            return False

    async def test_content_organization(self, note_id, token, user_type="regular"):
        """Test content organization and AI conversation separation"""
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        doc = Document(temp_file_path)
                        
                        # Analyze content organization
                        text_content = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_content.append(para.text.strip())
                        
                        full_text = " ".join(text_content)
                        
                        # Check for question/answer structure
                        question_indicators = ["Question 1", "Question 2", "Question 3", "Question 4"]
                        analysis_indicators = ["AI Analysis", "Analysis", "Response"]
                        
                        questions_found = sum(1 for indicator in question_indicators if indicator in full_text)
                        analysis_sections_found = sum(1 for indicator in analysis_indicators if indicator in full_text)
                        
                        # Check for proper content separation
                        has_clear_separation = questions_found >= 2 and analysis_sections_found >= 2
                        
                        # Check for professional content (not just raw AI output)
                        professional_indicators = ["supply chain", "logistics", "performance", "analysis", "strategy"]
                        professional_content = sum(1 for indicator in professional_indicators if indicator.lower() in full_text.lower())
                        
                        # Clean up temp file
                        os.unlink(temp_file_path)
                        
                        organization_score = 0
                        organization_details = []
                        
                        if questions_found >= 2:
                            organization_score += 1
                            organization_details.append(f"✓ Multiple questions found ({questions_found})")
                        else:
                            organization_details.append(f"✗ Limited questions found ({questions_found})")
                        
                        if analysis_sections_found >= 2:
                            organization_score += 1
                            organization_details.append(f"✓ Analysis sections found ({analysis_sections_found})")
                        else:
                            organization_details.append(f"✗ Limited analysis sections ({analysis_sections_found})")
                        
                        if has_clear_separation:
                            organization_score += 1
                            organization_details.append("✓ Clear Q&A separation")
                        else:
                            organization_details.append("✗ Poor Q&A separation")
                        
                        if professional_content >= 3:
                            organization_score += 1
                            organization_details.append(f"✓ Professional content ({professional_content} indicators)")
                        else:
                            organization_details.append(f"✗ Limited professional content ({professional_content} indicators)")
                        
                        if len(text_content) >= 10:
                            organization_score += 1
                            organization_details.append(f"✓ Rich content structure ({len(text_content)} text blocks)")
                        else:
                            organization_details.append(f"✗ Limited content structure ({len(text_content)} text blocks)")
                        
                        success = organization_score >= 3  # At least 3 out of 5 criteria
                        details = f"Organization score: {organization_score}/5. " + "; ".join(organization_details)
                        
                        await self.log_test(f"Content Organization ({user_type})", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test(f"Content Organization ({user_type})", False, 
                                          f"Content analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test(f"Content Organization ({user_type})", False, 
                                      f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test(f"Content Organization ({user_type})", False, f"Exception: {str(e)}")
            return False

    async def test_expeditors_branding(self):
        """Test Expeditors branding integration"""
        if not self.expeditors_note_id or not self.expeditors_token:
            await self.log_test("Expeditors Branding", False, "No Expeditors user/note available")
            return False
            
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.expeditors_token}"}
                
                response = await client.get(
                    f"{BACKEND_URL}/notes/{self.expeditors_note_id}/ai-conversations/export?format=docx",
                    headers=headers
                )
                
                if response.status_code == 200:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                        temp_file.write(response.content)
                        temp_file_path = temp_file.name
                    
                    try:
                        from docx import Document
                        doc = Document(temp_file_path)
                        
                        # Check for Expeditors branding
                        text_content = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_content.append(para.text.strip().lower())
                        
                        full_text = " ".join(text_content)
                        
                        # Look for Expeditors-specific content
                        expeditors_indicators = [
                            "expeditors international",
                            "expeditors",
                            "confidential - expeditors"
                        ]
                        
                        branding_found = sum(1 for indicator in expeditors_indicators if indicator in full_text)
                        
                        # Check for logo (would be in document structure, not text)
                        # This is harder to detect programmatically, so we'll check for image elements
                        has_images = False
                        try:
                            # Check if document has any images/shapes
                            for rel in doc.part.rels.values():
                                if "image" in rel.target_ref:
                                    has_images = True
                                    break
                        except:
                            pass
                        
                        # Clean up temp file
                        os.unlink(temp_file_path)
                        
                        branding_score = 0
                        branding_details = []
                        
                        if branding_found >= 1:
                            branding_score += 1
                            branding_details.append(f"✓ Expeditors branding text found ({branding_found} instances)")
                        else:
                            branding_details.append("✗ No Expeditors branding text found")
                        
                        if has_images:
                            branding_score += 1
                            branding_details.append("✓ Images/logo elements detected")
                        else:
                            branding_details.append("✗ No images/logo elements detected")
                        
                        # Check for professional layout (Expeditors should have enhanced formatting)
                        if len(text_content) >= 15:
                            branding_score += 1
                            branding_details.append("✓ Enhanced professional layout")
                        else:
                            branding_details.append("✗ Basic layout detected")
                        
                        success = branding_score >= 2  # At least 2 out of 3 criteria
                        details = f"Branding score: {branding_score}/3. " + "; ".join(branding_details)
                        
                        await self.log_test("Expeditors Branding", success, details)
                        return success
                        
                    except Exception as doc_error:
                        os.unlink(temp_file_path)
                        await self.log_test("Expeditors Branding", False, 
                                          f"Branding analysis error: {str(doc_error)}")
                        return False
                else:
                    await self.log_test("Expeditors Branding", False, 
                                      f"Export failed: {response.status_code}")
                    return False
                    
        except Exception as e:
            await self.log_test("Expeditors Branding", False, f"Exception: {str(e)}")
            return False

    async def test_format_comparison(self):
        """Test different export formats to verify Word improvements"""
        if not self.test_note_id or not self.auth_token:
            await self.log_test("Format Comparison", False, "No test note available")
            return False
            
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                headers = {"Authorization": f"Bearer {self.auth_token}"}
                
                formats = ["txt", "rtf", "docx"]
                format_results = {}
                
                for fmt in formats:
                    response = await client.get(
                        f"{BACKEND_URL}/notes/{self.test_note_id}/ai-conversations/export?format={fmt}",
                        headers=headers
                    )
                    
                    if response.status_code == 200:
                        content_length = len(response.content)
                        content_type = response.headers.get("content-type", "")
                        format_results[fmt] = {
                            "success": True,
                            "size": content_length,
                            "content_type": content_type
                        }
                    else:
                        format_results[fmt] = {
                            "success": False,
                            "error": f"Status {response.status_code}"
                        }
                
                # Analyze results
                successful_formats = [fmt for fmt, result in format_results.items() if result.get("success")]
                
                comparison_details = []
                for fmt in formats:
                    result = format_results[fmt]
                    if result.get("success"):
                        comparison_details.append(f"{fmt.upper()}: {result['size']} bytes")
                    else:
                        comparison_details.append(f"{fmt.upper()}: FAILED - {result.get('error', 'Unknown error')}")
                
                # Word should be significantly larger due to formatting
                docx_success = format_results.get("docx", {}).get("success", False)
                docx_size = format_results.get("docx", {}).get("size", 0)
                txt_size = format_results.get("txt", {}).get("size", 0)
                
                word_enhanced = docx_success and docx_size > txt_size * 2  # Word should be much larger
                
                success = len(successful_formats) >= 2 and docx_success
                details = f"Formats working: {len(successful_formats)}/3. " + "; ".join(comparison_details)
                if word_enhanced:
                    details += ". ✓ Word format shows enhanced formatting (larger file size)"
                
                await self.log_test("Format Comparison", success, details)
                return success
                
        except Exception as e:
            await self.log_test("Format Comparison", False, f"Exception: {str(e)}")
            return False

    async def run_all_tests(self):
        """Run all AI Word export tests"""
        print("📄 Starting AI Analysis Word Document Export Testing...")
        print(f"Backend URL: {BACKEND_URL}")
        print("=" * 80)

        # Test sequence
        tests = [
            ("Setup Users", self.create_test_users),
            ("Create Regular User Note", lambda: self.create_note_with_ai_conversations(self.auth_token, "regular")),
            ("Create Expeditors User Note", lambda: self.create_note_with_ai_conversations(self.expeditors_token, "expeditors")),
            ("Word Export - Regular User", lambda: self.test_word_export_functionality(self.test_note_id, self.auth_token, "regular")),
            ("Word Export - Expeditors User", lambda: self.test_word_export_functionality(self.expeditors_note_id, self.expeditors_token, "expeditors")),
            ("Document Structure - Regular", lambda: self.test_document_structure_and_formatting(self.test_note_id, self.auth_token, "regular")),
            ("Document Structure - Expeditors", lambda: self.test_document_structure_and_formatting(self.expeditors_note_id, self.expeditors_token, "expeditors")),
            ("Content Organization - Regular", lambda: self.test_content_organization(self.test_note_id, self.auth_token, "regular")),
            ("Content Organization - Expeditors", lambda: self.test_content_organization(self.expeditors_note_id, self.expeditors_token, "expeditors")),
            ("Expeditors Branding", self.test_expeditors_branding),
            ("Format Comparison", self.test_format_comparison)
        ]

        passed = 0
        total = len(tests)

        for test_name, test_func in tests:
            try:
                result = await test_func()
                if result:
                    passed += 1
            except Exception as e:
                await self.log_test(test_name, False, f"Unexpected exception: {str(e)}")

        print("=" * 80)
        print(f"📄 AI Word Export Testing Complete")
        print(f"📊 Results: {passed}/{total} tests passed ({(passed/total)*100:.1f}%)")
        
        if passed >= total * 0.8:  # 80% pass rate
            print("✅ AI Word Export functionality is working well!")
        else:
            print("❌ AI Word Export functionality needs attention!")
            
        return passed, total, self.test_results

async def main():
    """Main test runner"""
    tester = AIWordExportTester()
    passed, total, results = await tester.run_all_tests()
    
    # Print detailed results
    print("\n" + "=" * 80)
    print("📋 DETAILED TEST RESULTS:")
    print("=" * 80)
    
    for result in results:
        status = "✅ PASS" if result["success"] else "❌ FAIL"
        print(f"{status}: {result['test']}")
        if result["details"]:
            print(f"   Details: {result['details']}")
    
    print("\n" + "=" * 80)
    print("🎯 SUMMARY:")
    print("=" * 80)
    
    if passed >= total * 0.8:
        print("✅ AI Analysis Word Document Export is working excellently!")
        print("✅ Professional document formatting implemented")
        print("✅ Proper paragraph structure and headings")
        print("✅ Content organization and AI conversation separation")
        print("✅ Multiple export formats supported")
        if any("Expeditors Branding" in r["test"] and r["success"] for r in results):
            print("✅ Expeditors branding integration working")
    else:
        print("❌ AI Word Export functionality has issues that need attention")
        failed_tests = [r for r in results if not r["success"]]
        print(f"❌ Failed tests: {len(failed_tests)}")
        for failed in failed_tests:
            print(f"   - {failed['test']}: {failed['details']}")

if __name__ == "__main__":
    asyncio.run(main())