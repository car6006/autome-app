from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks, Query, Depends, status, Body
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import time
from pathlib import Path
import hashlib
import secrets
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
import httpx
import time
from datetime import datetime, timedelta, timezone
from openai import OpenAI

from store import NotesStore
from storage import store_file
from tasks import enqueue_transcription, enqueue_ocr, enqueue_email, enqueue_git_sync, enqueue_iisb_processing
from auth import (
    AuthService, User, UserCreate, UserLogin, UserResponse, UserProfileUpdate, 
    Token, get_current_user, get_current_user_optional
)

# Import new APIs
from upload_api import router as upload_router
from transcription_api import router as transcription_router
from webhooks import webhook_router

# Import Live Transcription System
from streaming_endpoints import streaming_router
from live_transcription import live_transcription_manager

# Phase 4: Production imports
from cloud_storage import storage_manager
from cache_manager import cache_manager
from monitoring import monitoring_service, monitor_endpoint
from rate_limiting import rate_limiter, quota_manager, check_rate_limit, check_user_quota
from webhooks import webhook_manager
from transcription_api import router as transcription_router
from worker_manager import worker_lifespan, get_pipeline_status

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Set up logging
logger = logging.getLogger(__name__)

# API Key validation on startup
def validate_api_keys():
    """Validate critical API keys on startup"""
    whisper_key = os.environ.get('WHISPER_API_KEY')
    gcv_key = os.environ.get('GCV_API_KEY')
    sendgrid_key = os.environ.get('SENDGRID_API_KEY')
    
    warnings = []
    if not whisper_key or whisper_key == 'your_openai_api_key_here':
        warnings.append("⚠️  WHISPER_API_KEY is missing or invalid - audio transcription will fail")
    if not gcv_key or gcv_key == 'your_google_vision_api_key_here':
        warnings.append("⚠️  GCV_API_KEY is missing or invalid - OCR processing will fail")
    if not sendgrid_key or sendgrid_key == 'your_sendgrid_api_key_here':
        warnings.append("⚠️  SENDGRID_API_KEY is missing or invalid - email delivery will fail")
    
    if warnings:
        print("\n" + "="*80)
        print("🚨 API KEY VALIDATION WARNINGS:")
        for warning in warnings:
            print(warning)
        print("="*80 + "\n")
    else:
        print("✅ All API keys appear to be configured")

# Validate keys on startup
validate_api_keys()

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI(title="AUTO-ME Productivity API", version="2.0.0", lifespan=worker_lifespan)

# 🔒 SECURITY MIDDLEWARE
# Trusted Host Middleware (prevent host header attacks)
app.add_middleware(
    TrustedHostMiddleware, 
    allowed_hosts=["*"],  # In production: set to specific domains only
)

# Security Headers Middleware
@app.middleware("http")
async def add_security_headers(request, call_next):
    response = await call_next(request)
    
    # Security headers to prevent common attacks
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Content-Security-Policy"] = "default-src 'self'; script-src 'self' 'unsafe-inline' 'unsafe-eval'; style-src 'self' 'unsafe-inline';"
    response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    
    # Prevent directory traversal and URL manipulation
    if "../" in str(request.url) or "..\\" in str(request.url):
        raise HTTPException(status_code=400, detail="Invalid URL format")
    
    return response

# Input Validation and Rate Limiting
@app.middleware("http")
async def validate_and_rate_limit(request, call_next):
    """Enhanced request validation and centralized rate limiting"""
    
    # Security validation - check for malicious patterns
    malicious_patterns = [
        "../", "\\", "cmd", "exec", "eval", "script", "javascript:",
        "data:", "vbscript:", "onload", "onerror", "onclick", "onmouseover",
        "<?php", "<%", "{{", "{%", "<%=", "#{", "${", "/*", "*/"
    ]
    
    # Check URL and query parameters
    url_path = str(request.url.path).lower()
    query_string = str(request.url.query).lower()
    
    # Skip security checks for legitimate API endpoints
    legitimate_endpoints = ['/api/', '/transcriptions/', '/metrics', '/auth/', '/notes', '/upload']
    is_legitimate_api = any(endpoint in url_path for endpoint in legitimate_endpoints)
    
    if not is_legitimate_api:
        for pattern in malicious_patterns:
            if pattern in url_path or pattern in query_string:
                logger.warning(f"🚨 Blocked malicious request: {pattern} in {request.url}")
                raise HTTPException(
                    status_code=400, 
                    detail="Request blocked for security reasons"
                )
    
    # Centralized rate limiting using proper rate_limiting.py utilities
    from rate_limiting import check_rate_limit, check_user_quota
    
    # Extract user information
    user_id = 'anonymous'
    user_tier = 'free'
    
    # Try to get user info from JWT token if present
    try:
        auth_header = request.headers.get('authorization')
        if auth_header and auth_header.startswith('Bearer '):
            from auth import decode_token
            token = auth_header.split(' ')[1]
            user_data = decode_token(token)
            if user_data:
                user_id = user_data.get('user_id', 'anonymous')
                # Could get user_tier from database if needed
        else:
            # Use IP for anonymous users
            user_id = f"ip_{request.client.host}"
    except Exception:
        # Fallback to IP-based identification
        user_id = f"ip_{request.client.host}"
    
    # Determine endpoint category for rate limiting
    endpoint_category = "general"
    if "/upload" in url_path:
        endpoint_category = "upload"
    elif "/transcription" in url_path or "/transcribe" in url_path:
        endpoint_category = "transcription"
    
    # Check rate limits
    rate_allowed, rate_status = await check_rate_limit(user_id, endpoint_category)
    if not rate_allowed:
        logger.warning(f"🚨 Rate limit exceeded for {user_id} on {endpoint_category}")
        raise HTTPException(
            status_code=429,
            detail={
                "error": "Rate limit exceeded",
                "limit_info": rate_status,
                "retry_after": 60
            }
        )
    
    # Check user quotas for authenticated users
    if user_id != 'anonymous' and not user_id.startswith('ip_'):
        quota_allowed, quota_status = await check_user_quota(user_id, user_tier, api_calls=1)
        if not quota_allowed:
            logger.warning(f"🚨 Quota exceeded for user {user_id}")
            raise HTTPException(
                status_code=429,
                detail={
                    "error": "Quota exceeded",
                    "quota_info": quota_status,
                    "violations": quota_status.get('violations', [])
                }
            )
    
    return await call_next(request)

# Global exception handler for enhanced security
@app.exception_handler(500)
async def internal_server_error_handler(request, exc):
    """Generic handler for internal server errors to prevent stack trace exposure"""
    logger.error(f"Internal server error: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error occurred"}
    )

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    """Catch-all exception handler for enhanced security"""
    logger.error(f"Unhandled exception: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Service temporarily unavailable"}
    )

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Models
class BatchReportRequest(BaseModel):
    note_ids: List[str]
    title: str = "Combined Analysis Report"
    format: str = "professional"  # Options: "professional", "txt", "rtf"

class NoteCreate(BaseModel):
    title: str
    kind: str  # "audio", "photo", or "text"
    text_content: Optional[str] = None

class EmailRequest(BaseModel):
    to: List[str]
    subject: str

class NoteResponse(BaseModel):
    id: str
    title: str
    kind: str
    status: str
    artifacts: Dict[str, Any] = {}
    metrics: Dict[str, Any] = {}
    created_at: datetime
    ready_at: Optional[datetime] = None
    user_id: Optional[str] = None

# Authentication endpoints
@api_router.post("/auth/verify-user")
async def verify_user(request: dict):
    """Verify if user exists for password reset"""
    email = request.get("email")
    if not email:
        raise HTTPException(status_code=400, detail="Email is required")
    
    # Clean and normalize email (trim whitespace and convert to lowercase)
    email = email.strip().lower()
    
    # Find user by email
    user = await AuthService.get_user_by_email(email)
    
    if user:
        return {"exists": True, "message": "User found"}
    else:
        raise HTTPException(status_code=404, detail="User not found")

@api_router.post("/auth/reset-password")
async def reset_password(request: dict):
    """Reset user password"""
    email = request.get("email")
    new_password = request.get("newPassword")
    
    if not email or not new_password:
        raise HTTPException(status_code=400, detail="Email and new password are required")
    
    if len(new_password) < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters long")
    
    # Find user by email
    user = await AuthService.get_user_by_email(email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Update password
    success = await AuthService.update_user_password(user["id"], new_password)
    
    if success:
        return {"message": "Password updated successfully"}
    else:
        raise HTTPException(status_code=500, detail="Failed to update password")

@api_router.post("/auth/validate-email")
async def validate_email_for_reset(request: dict):
    """Validate if email exists in the system for password reset"""
    try:
        email = request.get("email")
        if not email:
            raise HTTPException(status_code=400, detail="Email is required")
        
        email_exists = await AuthService.validate_email_exists(email)
        
        if email_exists:
            return {
                "message": "Email validated successfully",
                "email_exists": True
            }
        else:
            raise HTTPException(
                status_code=404, 
                detail="Email not found in our system"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Email validation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Email validation failed")

@api_router.post("/auth/reset-password")
async def reset_password(request: dict):
    """Reset password using email validation"""
    try:
        email = request.get("email")
        new_password = request.get("new_password")
        
        if not email or not new_password:
            raise HTTPException(
                status_code=400, 
                detail="Email and new password are required"
            )
        
        # Validate password strength
        if len(new_password) < 6:
            raise HTTPException(
                status_code=400, 
                detail="Password must be at least 6 characters long"
            )
        
        # Reset password
        success = await AuthService.reset_password_by_email(email, new_password)
        
        if success:
            return {
                "message": "Password reset successfully",
                "success": True
            }
        else:
            raise HTTPException(
                status_code=404, 
                detail="Email not found or password reset failed"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Password reset error: {str(e)}")
        raise HTTPException(status_code=500, detail="Password reset failed")

@api_router.post("/auth/register", response_model=Token)
async def register(user_data: UserCreate):
    """Register a new user"""
    user_id = await AuthService.create_user(user_data)
    user = await AuthService.get_user_by_id(user_id)
    
    # Create access token
    access_token = AuthService.create_access_token(data={"sub": user_id})
    
    return Token(
        access_token=access_token,
        token_type="bearer",
        user=UserResponse(**user)
    )

@api_router.post("/auth/login", response_model=Token)
async def login(user_data: UserLogin):
    """Login user"""
    user = await AuthService.authenticate_user(user_data.email, user_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password"
        )
    
    # Create access token
    access_token = AuthService.create_access_token(data={"sub": user["id"]})
    
    return Token(
        access_token=access_token,
        token_type="bearer",
        user=UserResponse(**user)
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_current_user_profile(current_user: dict = Depends(get_current_user)):
    """Get current user profile"""
    return UserResponse(**current_user)

@api_router.put("/auth/me", response_model=UserResponse)
async def update_profile(
    profile_data: UserProfileUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update user profile"""
    success = await AuthService.update_user_profile(current_user["id"], profile_data)
    if not success:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Failed to update profile"
        )
    
    # Get updated user
    updated_user = await AuthService.get_user_by_id(current_user["id"])
    return UserResponse(**updated_user)

# Core endpoints (updated to work with authentication)
@api_router.get("/")
async def root():
    return {"message": "AUTO-ME Productivity API v2.0", "status": "running"}

@api_router.post("/notes", response_model=Dict[str, str])
async def create_note(
    note: NoteCreate,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Create a new note (authenticated or anonymous)"""
    user_id = current_user["id"] if current_user else None
    note_id = await NotesStore.create(note.title, note.kind, user_id)
    
    # For text notes, immediately set the content and mark as ready
    if note.kind == "text" and note.text_content:
        artifacts = {"text": note.text_content}
        await NotesStore.set_artifacts(note_id, artifacts)
        await NotesStore.update_status(note_id, "ready")
        return {"id": note_id, "status": "ready"}
    
    return {"id": note_id, "status": "created"}


# Hidden IISB Analysis Feature (Expeditors only)
@api_router.post("/iisb/analyze", response_model=Dict[str, Any])
async def analyze_client_issues(
    request: Dict[str, str],
    current_user: dict = Depends(get_current_user)
):
    """Analyze client supply chain issues using IISB framework (HIDDEN - Expeditors only)"""
    # Check if user has @expeditors.com email
    if not current_user["email"].endswith("@expeditors.com"):
        raise HTTPException(
            status_code=404, 
            detail="Feature not found"
        )
    
    client_name = request.get("client_name", "")
    issues_text = request.get("issues_text", "")
    
    if not client_name or not issues_text:
        raise HTTPException(
            status_code=400,
            detail="Client name and issues text are required"
        )
    
    # Process IISB analysis
    from iisb_processor import iisb_processor
    result = await iisb_processor.process_iisb_input(client_name, issues_text)
    
    return result

@api_router.post("/notes/{note_id}/continue-to-iisb")
async def continue_to_iisb_analysis(
    note_id: str,
    request: Dict[str, str],
    current_user: dict = Depends(get_current_user)
):
    """IISB analysis (HIDDEN - Expeditors only)"""
    # Check if user has @expeditors.com email
    if not current_user["email"].endswith("@expeditors.com"):
        raise HTTPException(
            status_code=404, 
            detail="Feature not found"
        )
    
    # Verify note exists and belongs to user
    note = await NotesStore.get(note_id)
    if not note or note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Extract client name from note if available
    client_name = request.get("client_name", "")
    if not client_name and note.get("artifacts"):
        # Try to extract client name from note title or artifacts
        client_name = note.get("title", "").split("-")[0].strip() or "Client"
    
    return {
        "note_id": note_id,
        "client_name": client_name,
        "network_completed": note.get("status") == "ready",
        "ready_for_iisb": True,
        "message": "Ready to proceed with IISB analysis"
    }

@api_router.post("/notes/{note_id}/upload")
async def upload_media(
    note_id: str,
    background_tasks: BackgroundTasks,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    file: UploadFile = File(...)
):
    """Upload media file for a note"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    # Store the file
    file_content = await file.read()
    media_key = store_file(file_content, file.filename)
    
    # Update note with media key
    await NotesStore.update_media_key(note_id, media_key)
    
    # Queue processing based on note kind
    if note["kind"] == "audio":
        background_tasks.add_task(enqueue_transcription, note_id)
    elif note["kind"] == "photo":
        background_tasks.add_task(enqueue_ocr, note_id)

    
    return {"message": "File uploaded successfully", "status": "processing"}

@api_router.post("/upload-file")
async def upload_file_for_scan(
    background_tasks: BackgroundTasks,
    current_user: Optional[dict] = Depends(get_current_user_optional),
    file: UploadFile = File(...),
    title: str = Form("Uploaded Document")
):
    """Upload file directly for processing (scan or audio)"""
    
    # Validate file type
    image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.pdf'}
    audio_extensions = {'.mp3', '.wav', '.m4a', '.webm', '.ogg', '.mpeg'}
    allowed_extensions = image_extensions | audio_extensions
    
    file_ext = os.path.splitext(file.filename or '')[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Allowed: Images ({', '.join(image_extensions)}), Audio ({', '.join(audio_extensions)})"
        )
    
    # Determine note kind based on file type
    if file_ext in audio_extensions or file.content_type.startswith('audio/'):
        note_kind = "audio"
    else:
        note_kind = "photo"
    
    # Create note for the upload
    user_id = current_user["id"] if current_user else None
    note_id = await NotesStore.create(title, note_kind, user_id)
    
    # Store the file
    file_content = await file.read()
    media_key = store_file(file_content, file.filename)
    
    # Update note with media key
    await NotesStore.update_media_key(note_id, media_key)
    
    # Queue appropriate processing
    if note_kind == "audio":
        background_tasks.add_task(enqueue_transcription, note_id)
    else:
        background_tasks.add_task(enqueue_ocr, note_id)
    
    return {
        "id": note_id,
        "message": f"{'Audio' if note_kind == 'audio' else 'File'} uploaded successfully",
        "status": "processing",
        "filename": file.filename,
        "kind": note_kind
    }

@api_router.get("/notes/failed-count")
async def get_failed_notes_count(
    current_user: dict = Depends(get_current_user)
):
    """Get count of failed/stuck notes for the current user"""
    try:
        user_id = current_user["id"]
        
        failed_count = await db["notes"].count_documents({
            "$and": [
                {"user_id": user_id},
                {
                    "$or": [
                        {"status": {"$in": ["failed", "error", "stuck"]}},
                        {"artifacts.error": {"$exists": True}},
                        {
                            "$and": [
                                {"status": "processing"},
                                {"created_at": {"$lt": datetime.now(timezone.utc) - timedelta(hours=1)}}
                            ]
                        }
                    ]
                }
            ]
        })
        
        return {
            "failed_count": failed_count,
            "has_failed_notes": failed_count > 0
        }
        
    except Exception as e:
        logger.error(f"Failed to get failed notes count for user {current_user.get('id')}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to get failed notes count")

@api_router.post("/notes/cleanup-failed")
async def cleanup_failed_notes(
    current_user: dict = Depends(get_current_user)
):
    """Clean up failed, stuck, or error notes for the current user"""
    try:
        user_id = current_user["id"]
        
        # Define statuses and conditions that indicate failed/stuck notes
        cleanup_conditions = {
            "$and": [
                {"user_id": user_id},  # Only user's own notes
                {
                    "$or": [
                        {"status": {"$in": ["failed", "error", "stuck"]}},
                        {"artifacts.error": {"$exists": True}},
                        {
                            "$and": [
                                {"status": "processing"},
                                {"created_at": {"$lt": datetime.now(timezone.utc) - timedelta(hours=1)}}  # Processing for over 1 hour
                            ]
                        }
                    ]
                }
            ]
        }
        
        # First, get the notes that will be deleted for reporting
        notes_to_delete = await db["notes"].find(cleanup_conditions).to_list(length=100)
        
        # Delete the failed/stuck notes
        result = await db["notes"].delete_many(cleanup_conditions)
        
        # Categorize what was deleted
        deleted_by_status = {}
        for note in notes_to_delete:
            status = note.get("status", "unknown")
            if status not in deleted_by_status:
                deleted_by_status[status] = 0
            deleted_by_status[status] += 1
        
        logger.info(f"🧹 User {current_user.get('email')} cleaned up {result.deleted_count} failed notes: {deleted_by_status}")
        
        return {
            "message": f"Successfully cleaned up {result.deleted_count} failed/stuck notes",
            "deleted_count": result.deleted_count,
            "deleted_by_status": deleted_by_status,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Failed to cleanup notes for user {current_user.get('id')}: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to cleanup notes")

@api_router.get("/notes/{note_id}", response_model=NoteResponse)
async def get_note(
    note_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get a specific note (authentication required)"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    return NoteResponse(**note)

@api_router.get("/notes", response_model=List[NoteResponse])
async def list_notes(
    limit: int = Query(50, le=100),
    current_user: dict = Depends(get_current_user)
):
    """List user's notes (authentication required)"""
    user_id = current_user["id"]
    notes = await NotesStore.list_recent(limit, user_id)
    return [NoteResponse(**note) for note in notes]

@api_router.post("/notes/{note_id}/email")
async def send_note_email(
    note_id: str,
    email_req: EmailRequest,
    background_tasks: BackgroundTasks,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Send note content via email"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    background_tasks.add_task(enqueue_email, note_id, email_req.to, email_req.subject)
    return {"message": "Email queued for delivery"}

@api_router.delete("/notes/{note_id}")
async def delete_note(
    note_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a specific note (authentication required)"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to delete this note")
    
    # Delete from database
    await db["notes"].delete_one({"id": note_id})
    
    return {"message": "Note deleted successfully"}




@api_router.post("/user/professional-context")
async def update_professional_context(
    context_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Update user's professional context for personalized AI responses"""
    user_id = current_user["id"]
    
    # Extract professional context fields
    professional_updates = {
        "primary_industry": context_data.get("primary_industry", ""),
        "job_role": context_data.get("job_role", ""),
        "work_environment": context_data.get("work_environment", ""),
        "key_focus_areas": context_data.get("key_focus_areas", []),
        "content_types": context_data.get("content_types", []),
        "analysis_preferences": context_data.get("analysis_preferences", [])
    }
    
    # Update user profile
    await db["users"].update_one(
        {"id": user_id},
        {"$set": {f"profile.{key}": value for key, value in professional_updates.items()}}
    )
    
    return {"message": "Professional context updated successfully", "context": professional_updates}

@api_router.get("/health")
@monitor_endpoint("health_check")
async def health_check():
    """Enhanced health check endpoint with Phase 4 monitoring"""
    try:
        # Check database connection
        await db["users"].find_one({})
        database_health = "healthy"
        
        # Check pipeline status
        try:
            from worker_manager import get_pipeline_status
            pipeline_status = await get_pipeline_status()
            pipeline_health = "healthy" if pipeline_status.get("worker", {}).get("running") else "degraded"
        except Exception as e:
            pipeline_health = "unknown"
            pipeline_status = {"error": str(e)}
        
        # Phase 4: Check production services
        services_health = {
            "database": database_health,
            "api": "running",
            "pipeline": pipeline_health
        }
        
        # Check cache status
        try:
            cache_stats = await cache_manager.get_cache_stats()
            services_health["cache"] = "healthy" if cache_stats.get("enabled") else "disabled"
        except Exception as e:
            services_health["cache"] = "error"
        
        # Check storage status
        try:
            storage_stats = storage_manager.get_usage_stats()
            services_health["storage"] = "healthy"
        except Exception as e:
            services_health["storage"] = "error"
        
        # Check webhook manager
        try:
            webhook_stats = await webhook_manager.get_delivery_stats()
            services_health["webhooks"] = "healthy" if webhook_stats.get("enabled") else "disabled"
        except Exception as e:
            services_health["webhooks"] = "error"
        
        # Get comprehensive metrics
        try:
            metrics = await monitoring_service.get_comprehensive_metrics()
        except Exception as e:
            logger.warning(f"Failed to get comprehensive metrics: {e}")
            metrics = {"error": "metrics_unavailable"}
        
        # Determine overall health
        critical_services = ["database", "api", "pipeline"]
        unhealthy_critical = [s for s in critical_services if services_health.get(s) not in ["healthy", "running"]]
        
        if unhealthy_critical:
            overall_status = "unhealthy"
        elif any(status == "degraded" for status in services_health.values()):
            overall_status = "degraded"
        else:
            overall_status = "healthy"
        
        return {
            "status": overall_status,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "services": services_health,
            "pipeline": pipeline_status,
            "metrics": metrics,
            "version": "Phase 4 Production",
            "uptime_hours": (datetime.now(timezone.utc) - datetime.fromtimestamp(time.time() - 3600, timezone.utc)).total_seconds() / 3600
        }
        
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(
            status_code=503,
            detail={
                "status": "unhealthy", 
                "timestamp": datetime.now(timezone.utc).isoformat(),
                "error": "Health check failed",
                "details": str(e)
            }
        )

@api_router.get("/system-metrics")
@monitor_endpoint("system_metrics")
async def get_system_metrics(current_user: Optional[dict] = Depends(get_current_user_optional)):
    """Get comprehensive system metrics (Phase 4)"""
    
    # Check if user has access to metrics (admin only in production)
    if not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    # In production, restrict to admin users
    user_role = current_user.get("role", "user")
    if user_role not in ["admin", "enterprise"]:
        # Return limited metrics for regular users
        try:
            cache_stats = await cache_manager.get_cache_stats()
            storage_stats = storage_manager.get_usage_stats()
            
            return {
                "user_metrics": {
                    "user_id": current_user["id"],
                    "cache_enabled": cache_stats.get("enabled", False),
                    "storage_usage_bytes": storage_stats.get("bytes_stored", 0),
                    "files_stored": storage_stats.get("files_stored", 0)
                },
                "access_level": "limited",
                "timestamp": datetime.now(timezone.utc).isoformat()
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail="Failed to retrieve user metrics")
    
    # Full metrics for admin users
    try:
        comprehensive_metrics = await monitoring_service.get_comprehensive_metrics()
        
        # Add Phase 4 service metrics
        phase4_metrics = {
            "cache": await cache_manager.get_cache_stats(),
            "storage": storage_manager.get_usage_stats(),
            "rate_limiting": rate_limiter.get_limits_status("system"),
            "webhooks": await webhook_manager.get_delivery_stats()
        }
        
        return {
            **comprehensive_metrics,
            "phase4_services": phase4_metrics,
            "access_level": "full",
            "generated_for": current_user["id"],
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Failed to get system metrics: {e}")
        raise HTTPException(
            status_code=500,
            detail={
                "error": "Failed to retrieve system metrics",
                "details": str(e)
            }
        )

# ================================
# ARCHIVE MANAGEMENT ENDPOINTS
# ================================

@api_router.get("/admin/archive/status")
async def get_archive_status(
    current_user: dict = Depends(get_current_user)
):
    """Get current archive configuration and statistics"""
    try:
        import sys
        import os
        sys.path.append('/app/backend')
        from archive_manager import ArchiveManager
        
        archive_manager = ArchiveManager()
        
        # Get current configuration
        config = {
            "archive_days": archive_manager.ARCHIVE_DAYS,
            "storage_paths": archive_manager.STORAGE_PATHS,
            "archive_patterns": archive_manager.ARCHIVE_PATTERNS,
            "delete_patterns": archive_manager.DELETE_PATTERNS
        }
        
        # Run dry run to get statistics
        result = await archive_manager.run_archive_process(dry_run=True)
        
        return {
            "config": config,
            "statistics": result,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Failed to get archive status: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get archive status: {str(e)}")

@api_router.post("/admin/archive/run")
async def run_archive_process(
    dry_run: bool = Query(False, description="Run in dry-run mode without deleting files"),
    current_user: dict = Depends(get_current_user)
):
    """Run the archive process to clean up old files"""
    try:
        import sys
        import os
        sys.path.append('/app/backend')
        from archive_manager import ArchiveManager
        
        logger.info(f"Archive process started by user: {current_user.get('email', 'unknown')}")
        
        archive_manager = ArchiveManager()
        result = await archive_manager.run_archive_process(dry_run=dry_run)
        
        # Log the archive action for audit
        if not dry_run and result.get('success'):
            logger.info(f"🏛️ Archive completed by {current_user.get('email')}: "
                       f"freed {result.get('disk_space_freed_formatted', '0B')} "
                       f"({result.get('total_processed', 0)} files)")
        
        return result
        
    except Exception as e:
        logger.error(f"Archive process failed: {e}")
        raise HTTPException(status_code=500, detail=f"Archive process failed: {str(e)}")

@api_router.post("/admin/archive/configure")
async def configure_archive_settings(
    archive_days: int = Body(..., description="Number of days after which files should be archived"),
    current_user: dict = Depends(get_current_user)
):
    """Configure archive settings"""
    try:
        if archive_days < 1:
            raise HTTPException(status_code=400, detail="Archive days must be at least 1")
        
        if archive_days > 365:
            raise HTTPException(status_code=400, detail="Archive days cannot exceed 365")
        
        # Update environment variable (would need restart to take effect)
        os.environ['ARCHIVE_DAYS'] = str(archive_days)
        
        logger.info(f"Archive settings updated by {current_user.get('email')}: {archive_days} days")
        
        return {
            "success": True,
            "archive_days": archive_days,
            "message": "Archive settings updated. Restart required for changes to take full effect.",
            "updated_by": current_user.get('email'),
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Failed to configure archive settings: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to configure archive settings: {str(e)}")

@api_router.get("/user/professional-context")
async def get_professional_context(
    current_user: dict = Depends(get_current_user)
):
    """Get user's current professional context"""
    user_id = current_user["id"]
    user = await db["users"].find_one({"id": user_id})
    
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    
    profile = user.get("profile", {})
    
    return {
        "primary_industry": profile.get("primary_industry", ""),
        "job_role": profile.get("job_role", ""),
        "work_environment": profile.get("work_environment", ""),
        "key_focus_areas": profile.get("key_focus_areas", []),
        "content_types": profile.get("content_types", []),
        "analysis_preferences": profile.get("analysis_preferences", [])
    }

from ai_context_processor import ai_context_processor
from enhanced_providers import transcribe_audio, generate_ai_analysis

@api_router.post("/batch-report/ai-chat")
async def batch_report_ai_chat(
    request: dict,
    current_user: dict = Depends(get_current_user)
):
    """Handle AI chat for batch reports without requiring database note lookup"""
    try:
        batch_content = request.get("content", "")
        question = request.get("question", "")
        
        if not batch_content or not question:
            raise HTTPException(status_code=400, detail="Both 'content' and 'question' are required")
        
        # Prepare context for AI
        context = f"""
        BATCH REPORT CONTENT:
        {batch_content}
        
        USER QUESTION: {question}
        
        Please provide a helpful response based on the batch report content above.
        """
        
        # Get AI response using OpenAI
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("WHISPER_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        async with httpx.AsyncClient(timeout=45) as client:
            openai_response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "user", "content": context}
                    ],
                    "max_tokens": 1200,
                    "temperature": 0.3
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            openai_response.raise_for_status()
            
            ai_analysis = openai_response.json()
            response = ai_analysis["choices"][0]["message"]["content"]
        
        return {
            "response": response,
            "question": question,
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "user_id": current_user["id"]
        }
        
    except Exception as e:
        print(f"Error in batch report AI chat: {e}")
        raise HTTPException(status_code=500, detail="Failed to process AI chat request")

@api_router.post("/notes/{note_id}/ai-chat")
async def ai_chat_with_note(
    note_id: str,
    request: dict,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """AI conversational agent for note content analysis with dynamic profession-based context"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    content = artifacts.get("transcript") or artifacts.get("text", "")
    
    if not content:
        raise HTTPException(status_code=400, detail="No content available for analysis")
    
    user_question = request.get("question", "").strip()
    if not user_question:
        raise HTTPException(status_code=400, detail="Please provide a question")
    
    try:
        # Use OpenAI with dynamic profession-based context
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("WHISPER_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        # Get user profile for context
        user_profile = current_user.get("profile", {}) if current_user else {}
        
        # Generate dynamic prompt based on user's profession and question
        if "meeting" in user_question.lower() or "minutes" in user_question.lower():
            analysis_type = "meeting_minutes"
        elif "insight" in user_question.lower() or "analysis" in user_question.lower():
            analysis_type = "insights"
        elif "summary" in user_question.lower() or "summarize" in user_question.lower():
            analysis_type = "summary"
        else:
            analysis_type = "general"
        
        # Create dynamic prompt with profession context
        base_prompt = ai_context_processor.generate_dynamic_prompt(
            content=content,
            user_profile=user_profile,
            analysis_type=analysis_type
        )
        
        # Add the specific user question
        full_prompt = f"""
        {base_prompt}
        
        User's specific question: {user_question}
        
        Provide a comprehensive, profession-specific response that directly addresses their question.
        """
        
        async with httpx.AsyncClient(timeout=45) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "user", "content": full_prompt}
                    ],
                    "max_tokens": 1200,
                    "temperature": 0.3
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            response.raise_for_status()
            
            ai_analysis = response.json()
            ai_response = ai_analysis["choices"][0]["message"]["content"]
            
            # Store the conversation in artifacts for reference
            conversations = artifacts.get("ai_conversations", [])
            conversations.append({
                "question": user_question,
                "response": ai_response,
                "context_used": ai_context_processor.get_context_summary(user_profile),
                "timestamp": datetime.now(timezone.utc).isoformat()
            })
            
            updated_artifacts = {**artifacts, "ai_conversations": conversations}
            await NotesStore.set_artifacts(note_id, updated_artifacts)
            
            return {
                "response": ai_response,
                "question": user_question,
                "note_title": note["title"],
                "context_summary": ai_context_processor.get_context_summary(user_profile),
                "timestamp": datetime.now(timezone.utc).isoformat()
            }
    
    except Exception as e:
        logger.error(f"AI chat processing error: {str(e)}")
        raise HTTPException(status_code=500, detail="AI chat service temporarily unavailable")

@api_router.post("/notes/{note_id}/generate-meeting-minutes")
async def generate_meeting_minutes(
    note_id: str,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Generate structured meeting minutes from AI conversations"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    conversations = artifacts.get("ai_conversations", [])
    transcript = artifacts.get("transcript") or artifacts.get("text", "")
    
    if not conversations and not transcript:
        raise HTTPException(status_code=400, detail="No content available for meeting minutes generation")
    
    try:
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("WHISPER_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="Service not configured")
        
        # Check if user is from Expeditors
        is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
        
        # Combine all available content
        all_content = []
        if transcript:
            all_content.append(transcript)
        
        for conv in conversations:
            response = conv.get("response", "")
            if response:
                all_content.append(response)
        
        combined_content = "\n\n".join(all_content)
        
        # Get user profile for professional context
        user_profile = current_user.get("profile", {}) if current_user else {}
        
        # Generate professional context-aware meeting minutes prompt
        meeting_minutes_prompt = ai_context_processor.generate_dynamic_prompt(
            content=combined_content,
            user_profile=user_profile,
            analysis_type="meeting_minutes"
        )
        
        # Enhanced meeting minutes prompt with professional context
        company_context = "Expeditors International (a global logistics and freight forwarding company)" if is_expeditors_user else "the organization"
        
        prompt = f"""
        {meeting_minutes_prompt}

        ADDITIONAL CONTEXT: You are creating formal meeting minutes for {company_context}.

        FORMATTING INSTRUCTIONS:
        Create comprehensive meeting minutes following this EXACT professional format with NO bold, NO markdown symbols:

        Date: [Extract or use current date]
        Meeting Type: [Determine meeting type from context]

        EXECUTIVE SUMMARY
        Write a brief paragraph summarizing the meeting's purpose, key topics discussed, and overall focus. Emphasize collaborative decision-making and strategic planning.

        OPENING REMARKS AND OBJECTIVES
        Describe how the meeting commenced, welcoming remarks, agenda setting, and primary objectives. Focus on the meeting's purpose and expected outcomes.

        KEY DISCUSSION TOPICS

        [Create appropriate section headers based on content, such as:]
        Performance Review and Analysis
        Strategic Planning and Resource Allocation  
        Customer Engagement and Service Delivery
        Operational Efficiency and Process Improvement
        Team Development and Communication
        Financial Performance Review
        [Add other relevant sections based on actual content]

        For each section, provide detailed narrative explaining:
        - What was discussed
        - Key points raised by participants  
        - Concerns or challenges identified
        - Solutions proposed
        - Decisions made

        DECISIONS AND RESOLUTIONS
        Summarize key decisions made during the meeting and any formal resolutions agreed upon.

        ACTION ITEMS AND NEXT STEPS
        Clearly defined action items with responsibilities assigned. Format as numbered list with specific, actionable descriptions.

        CONCLUSION  
        Professional summary emphasizing successful outcomes, renewed commitments, and next steps for implementation.

        IMPORTANT: 
        - Use professional business language appropriate for corporate meetings
        - Write in narrative form with detailed explanations
        - Include specific details mentioned in the transcript
        - Structure content logically with clear section divisions
        - NO bullet points in main content (only in action items)
        - Focus on business outcomes and strategic initiatives
        """
        
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "user", "content": prompt}
                    ],
                    "max_tokens": 2000,
                    "temperature": 0.3
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            response.raise_for_status()
            
            ai_analysis = response.json()
            meeting_minutes = ai_analysis["choices"][0]["message"]["content"]
            
            # Store the meeting minutes in artifacts
            updated_artifacts = {**artifacts, "meeting_minutes": meeting_minutes}
            await NotesStore.set_artifacts(note_id, updated_artifacts)
            
            return {
                "meeting_minutes": meeting_minutes,
                "note_title": note["title"],
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "note_id": note_id,
                "is_expeditors": is_expeditors_user
            }
    
    except Exception as e:
        logger.error(f"Meeting minutes generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Meeting minutes generation temporarily unavailable")

@api_router.post("/notes/{note_id}/generate-action-items")
async def generate_action_items(
    note_id: str,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Generate structured action items table from meeting content"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    transcript = artifacts.get("transcript", "")
    
    if not transcript:
        raise HTTPException(status_code=400, detail="No content available for action items generation")
    
    try:
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            raise HTTPException(status_code=503, detail="AI service configuration error")
        
        # Generate action items in clean, structured format
        prompt = f"""
        Based on the following meeting transcript, extract and create comprehensive action items in a professional format.

        TRANSCRIPT:
        {transcript}

        INSTRUCTIONS:
        Create structured action items with the following format:

        ACTION ITEMS FOR: ZA DM MEETING 5 SEPTEMBER 2025

        1. Schedule weekly meetings for all four participants to discuss ongoing topics and ensure alignment.

        2. Monitor and manage annual leave usage to prevent carryover beyond five days, ensuring compliance with legal requirements.

        3. Review and address any discrepancies in work-from-home reporting to ensure productivity is maintained.

        4. Check with Elaine and Garth regarding the dashboard for monitoring work-from-home days and ensure proper logging of attendance.

        5. Prepare for budget discussions by analyzing growth expectations and aligning team goals with corporate directives.

        REQUIREMENTS:
        - Extract ALL actionable items, decisions, and commitments mentioned
        - Write clear, specific action descriptions in complete sentences
        - Number each item sequentially (1, 2, 3, etc.)
        - Use professional business language
        - Each action item should be a complete paragraph
        - NO pipe characters or table formatting
        - NO markdown formatting, just clean text
        - Focus on concrete, measurable actions
        - Include both explicit action items and implied commitments

        Provide clean, readable text that can be easily copied into documents.
        """
        
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "user", "content": prompt}
                    ],
                    "max_tokens": 1500,
                    "temperature": 0.2
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            response.raise_for_status()
            
            ai_analysis = response.json()
            action_items_table = ai_analysis["choices"][0]["message"]["content"]
            
            # Store the action items in artifacts
            updated_artifacts = {**artifacts, "action_items": action_items_table}
            await NotesStore.set_artifacts(note_id, updated_artifacts)
            
            return {
                "action_items": action_items_table,
                "note_title": note["title"],
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "note_id": note_id
            }
    
    except Exception as e:
        logger.error(f"Action items generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Action items generation temporarily unavailable")

@api_router.get("/notes/{note_id}/action-items/export")
async def export_action_items(
    note_id: str,
    format: str = Query("txt", regex="^(txt|docx|rtf)$"),
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Export action items in clean formats (TXT, DOC, RTF)"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    action_items = artifacts.get("action_items", "")
    
    if not action_items:
        raise HTTPException(status_code=400, detail="No action items available. Generate action items first.")
    
    # Clean title for filename
    clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
    
    if format == "txt":
        # Clean TXT format with proper spacing
        content = f"ACTION ITEMS\n"
        content += "=" * 12 + "\n\n"
        content += f"Meeting: {note['title']}\n"
        content += f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}\n\n"
        
        # Clean up the action items text - remove pipe characters and improve formatting
        clean_action_items = action_items.replace("|", "").replace("  ", " ")
        
        # Split into lines and reformat
        lines = clean_action_items.split('\n')
        formatted_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('No.') and not line.startswith('---'):
                # If line starts with a number, it's an action item
                if line[0].isdigit():
                    formatted_lines.append(f"\n{line}")
                else:
                    # Continuation of previous item
                    formatted_lines.append(line)
        
        content += '\n'.join(formatted_lines)
        content += "\n\n" + "=" * 50 + "\n"
        content += "Ready to copy and customize for your meeting minutes\n"
        
        filename = f"{clean_title}_Action_Items.txt"
        
        return Response(
            content=content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "rtf":
        # RTF format with proper formatting
        filename = f"{clean_title}_Action_Items.rtf"
        
        # RTF Header
        rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
        
        # Title
        rtf_content += f"\\b\\fs32 ACTION ITEMS\\b0\\fs24\\par\\par "
        rtf_content += f"\\b Meeting:\\b0 {note['title']}\\par "
        rtf_content += f"\\b Generated:\\b0 {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}\\par\\par "
        
        # Clean and format action items
        clean_action_items = action_items.replace("|", "").replace("  ", " ")
        lines = clean_action_items.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('No.') and not line.startswith('---'):
                if line[0].isdigit():
                    rtf_content += f"\\par\\b {line.split('.')[0]}.\\b0 {'.'.join(line.split('.')[1:]).strip()}\\par "
                else:
                    rtf_content += f"{line} "
        
        rtf_content += "\\par\\par }"
        
        return Response(
            content=rtf_content,
            media_type="application/rtf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "docx":
        # Word document format
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            # Title
            title = doc.add_heading('ACTION ITEMS', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Meeting info
            doc.add_paragraph(f"Meeting: {note['title']}")
            doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y at %I:%M %p UTC')}")
            doc.add_paragraph("")  # Empty line
            
            # Clean and format action items
            clean_action_items = action_items.replace("|", "").replace("  ", " ")
            lines = clean_action_items.split('\n')
            
            for line in lines:
                line = line.strip()
                if line and not line.startswith('No.') and not line.startswith('---'):
                    if line[0].isdigit():
                        p = doc.add_paragraph()
                        # Add number in bold
                        num_part = line.split('.')[0] + '.'
                        desc_part = '.'.join(line.split('.')[1:]).strip()
                        
                        run = p.add_run(num_part)
                        run.bold = True
                        p.add_run(f" {desc_part}")
                    else:
                        # Continuation line
                        doc.add_paragraph(line, style='List Continue')
            
            # Add footer
            doc.add_paragraph("")  # Empty line
            footer_p = doc.add_paragraph("Ready to copy and customize for your meeting minutes")
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            filename = f"{clean_title}_Action_Items.docx"
            
            return Response(
                content=buffer.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
            )
            
        except ImportError:
            raise HTTPException(status_code=503, detail="DOCX export temporarily unavailable")

@api_router.get("/notes/{note_id}/ai-conversations/export")
async def export_ai_conversations(
    note_id: str,
    format: str = Query("pdf", regex="^(pdf|docx|txt|rtf)$"),
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Export AI conversations/analysis to PDF, Word DOC, TXT, or RTF format"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    
    # Try to get AI conversations first, then fall back to meeting minutes
    ai_conversations = artifacts.get("ai_conversations", [])
    meeting_minutes = artifacts.get("meeting_minutes", "")
    
    if not ai_conversations and not meeting_minutes:
        raise HTTPException(status_code=400, detail="No AI analysis or meeting minutes found. Please generate AI analysis first.")
    
    # Check if user is from Expeditors for branding
    is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
    
    # Clean content function to remove all AI formatting
    def clean_content(text):
        if not text:
            return text
        # Remove markdown formatting
        cleaned = text.replace("**", "").replace("***", "").replace("###", "").replace("##", "").replace("#", "").replace("*", "").replace("_", "")
        # Remove bullet points and section formatting
        lines = cleaned.split('\n')
        clean_lines = []
        for line in lines:
            line = line.strip()
            if line:
                line = line.lstrip('•-*1234567890. ')
                # Remove common AI section headers but keep content
                if not any(header in line.upper() for header in ['ATTENDEES:', 'APOLOGIES:', 'MEETING MINUTES:', 'ACTION ITEMS:', 'KEY INSIGHTS:', 'ASSESSMENTS:', 'RISK ASSESSMENT:', 'NEXT STEPS:']):
                    if line:
                        clean_lines.append(line)
        return '\n'.join(clean_lines)
    
    # Prepare content based on what's available
    content_title = note['title']
    export_content = ""
    
    if ai_conversations:
        export_content = f"AI Content Analysis for: {content_title}\n\n"
        for i, conv in enumerate(ai_conversations, 1):
            question = conv.get("question", "")
            response = conv.get("response", "")
            timestamp = conv.get("timestamp", "")
            
            export_content += f"Question {i}: {question}\n\n"
            export_content += f"AI Analysis:\n{clean_content(response)}\n\n"
            if i < len(ai_conversations):
                export_content += "=" * 50 + "\n\n"
    elif meeting_minutes:
        export_content = f"Meeting Minutes for: {content_title}\n\n{clean_content(meeting_minutes)}"
    
    if format == "txt":
        # Plain text format
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"AI_Analysis_{clean_title}.txt"
        
        return Response(
            content=export_content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "rtf":
        # RTF format
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"AI_Analysis_{clean_title}.rtf"
        
        # RTF header
        rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
        
        # Convert content to RTF
        rtf_body = export_content.replace("\\", "\\\\").replace("\n", "\\par ")
        rtf_content += rtf_body + "}"
        
        return Response(
            content=rtf_content,
            media_type="application/rtf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "pdf":
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import Color, black
        from io import BytesIO
        import os
        import re
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        # Create custom styles for professional AI analysis
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'AI Report Title',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=18,
            spaceBefore=0,
            alignment=1,  # Center
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        # Section heading style  
        section_heading_style = ParagraphStyle(
            'AI Section Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        # Sub-heading style for content within AI responses
        sub_heading_style = ParagraphStyle(
            'AI Sub Heading',
            parent=styles['Normal'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=black
        )
        
        # Body text style with improved spacing
        body_style = ParagraphStyle(
            'AI Body Text',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,  # Increased spacing between paragraphs
            spaceBefore=3,  # Small space before paragraphs
            leftIndent=0,
            rightIndent=0,
            fontName='Helvetica',
            leading=13.2  # 1.2 line spacing
        )
        
        # Question style
        question_style = ParagraphStyle(
            'Question Text',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=3,
            leftIndent=0,
            rightIndent=0,
            fontName='Helvetica',
            leading=13.2
        )
        
        # Function to format AI content for PDF with proper paragraph structure
        def format_ai_content_pdf(content):
            if not content:
                return []
            
            # Split by double newlines first to identify major sections
            sections = re.split(r'\n\s*\n', content)
            formatted_paragraphs = []
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Check if this looks like a heading (short line, possibly with special formatting)
                lines = section.split('\n')
                if len(lines) == 1 and len(section) < 100 and not section.endswith('.'):
                    # Likely a heading
                    heading_text = section.replace('**', '').replace('*', '').replace('#', '').strip()
                    if heading_text:
                        formatted_paragraphs.append(('heading', heading_text))
                else:
                    # Process as content section
                    for line in lines:
                        line = line.strip()
                        if line:
                            # Clean formatting but preserve structure
                            clean_line = line.replace('**', '').replace('*', '').replace('#', '').strip()
                            # Remove bullet points and numbering for cleaner look
                            if clean_line.startswith(('•', '-', '*')):
                                clean_line = clean_line[1:].strip()
                            elif re.match(r'^\d+\.', clean_line):
                                clean_line = re.sub(r'^\d+\.\s*', '', clean_line)
                            
                            if clean_line:
                                formatted_paragraphs.append(('paragraph', clean_line))
            
            return formatted_paragraphs
        
        story = []
        
        # Add logo if Expeditors user
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo = Image(logo_path, width=2*inch, height=0.8*inch)
                    logo.hAlign = 'CENTER'
                    story.append(logo)
                    story.append(Spacer(1, 20))
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title
        if is_expeditors_user:
            story.append(Paragraph("EXPEDITORS INTERNATIONAL", title_style))
            story.append(Paragraph("AI Content Analysis Report", section_heading_style))
            story.append(Paragraph(f"{content_title}", body_style))
        else:
            story.append(Paragraph("AI Content Analysis Report", title_style))
            story.append(Paragraph(f"{content_title}", section_heading_style))
        
        story.append(Spacer(1, 20))
        
        # Content
        if ai_conversations:
            for i, conv in enumerate(ai_conversations, 1):
                question = conv.get("question", "")
                response = conv.get("response", "")
                
                # Question section
                story.append(Paragraph(f"Question {i}", section_heading_style))
                story.append(Paragraph(question, question_style))
                
                # Small spacer between question and response
                story.append(Spacer(1, 6))
                
                # Response section
                story.append(Paragraph("Analysis & Response", section_heading_style))
                
                # Format the AI response with proper paragraphs
                formatted_content = format_ai_content_pdf(response)
                for content_type, content_text in formatted_content:
                    if content_type == 'heading':
                        # Sub-heading within the response
                        story.append(Paragraph(content_text, sub_heading_style))
                    else:
                        # Regular paragraph
                        story.append(Paragraph(content_text, body_style))
                
                # Add spacing between conversations
                if i < len(ai_conversations):
                    story.append(Spacer(1, 20))
        
        elif meeting_minutes:
            # Format meeting minutes with proper structure
            story.append(Paragraph("Meeting Minutes", section_heading_style))
            formatted_content = format_ai_content_pdf(meeting_minutes)
            for content_type, content_text in formatted_content:
                if content_type == 'heading':
                    story.append(Paragraph(content_text, sub_heading_style))
                else:
                    story.append(Paragraph(content_text, body_style))
        
        # Footer
        if is_expeditors_user:
            story.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer', 
                parent=styles['Normal'], 
                fontSize=8, 
                alignment=1, 
                textColor=Color(0.5, 0.5, 0.5),
                fontName='Helvetica-Oblique'
            )
            story.append(Paragraph("Confidential - Expeditors International", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Create filename
        filename_base = note['title'][:30].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"AI_Analysis_{filename_base}.pdf"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "docx":
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from docx.enum.style import WD_STYLE_TYPE
        from io import BytesIO
        import os
        import re
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define custom styles for better formatting
        styles = doc.styles
        
        # Main heading style
        if 'AI Report Title' not in [s.name for s in styles]:
            title_style = styles.add_style('AI Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(18)
        
        # Section heading style
        if 'AI Section Heading' not in [s.name for s in styles]:
            section_style = styles.add_style('AI Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'AI Body Text' not in [s.name for s in styles]:
            body_style = styles.add_style('AI Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(12)  # Increased from 6 to 12 for better spacing
            body_style.paragraph_format.space_before = Pt(3)  # Added space before paragraphs
            body_style.paragraph_format.line_spacing = 1.15
        
        # Function to format AI content with proper paragraph structure
        def format_ai_content(content):
            if not content:
                return []
            
            # Split by double newlines first to identify major sections
            sections = re.split(r'\n\s*\n', content)
            formatted_paragraphs = []
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Check if this looks like a heading (short line, possibly with special formatting)
                lines = section.split('\n')
                if len(lines) == 1 and len(section) < 100 and not section.endswith('.'):
                    # Likely a heading
                    heading_text = section.replace('**', '').replace('*', '').replace('#', '').strip()
                    if heading_text:
                        formatted_paragraphs.append(('heading', heading_text))
                else:
                    # Process as content section
                    for line in lines:
                        line = line.strip()
                        if line:
                            # Clean formatting but preserve structure
                            clean_line = line.replace('**', '').replace('*', '').replace('#', '').strip()
                            # Remove bullet points and numbering for cleaner look
                            if clean_line.startswith(('•', '-', '*')):
                                clean_line = clean_line[1:].strip()
                            elif re.match(r'^\d+\.', clean_line):
                                clean_line = re.sub(r'^\d+\.\s*', '', clean_line)
                            
                            if clean_line:
                                formatted_paragraphs.append(('paragraph', clean_line))
            
            return formatted_paragraphs
        
        # Add logo if Expeditors user
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
                    run.add_picture(logo_path, width=Inches(3))
                    doc.add_paragraph()  # Spacer
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title
        if is_expeditors_user:
            company_title = doc.add_paragraph('EXPEDITORS INTERNATIONAL', style='AI Report Title')
            analysis_title = doc.add_paragraph(f'AI Content Analysis Report', style='AI Section Heading')
            analysis_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_title = doc.add_paragraph(f'{content_title}', style='AI Body Text')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Explicitly set spacing for body text
            doc_title.paragraph_format.space_after = Pt(12)
            doc_title.paragraph_format.space_before = Pt(3)
        else:
            main_title = doc.add_paragraph(f'AI Content Analysis Report', style='AI Report Title')
            doc_title = doc.add_paragraph(f'{content_title}', style='AI Section Heading')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a horizontal line separator
        doc.add_paragraph()
        
        # Content
        if ai_conversations:
            for i, conv in enumerate(ai_conversations, 1):
                question = conv.get("question", "")
                response = conv.get("response", "")
                
                # Question section
                question_heading = doc.add_paragraph(f'Question {i}', style='AI Section Heading')
                question_para = doc.add_paragraph(question, style='AI Body Text')
                # Explicitly set spacing for body text
                question_para.paragraph_format.space_after = Pt(12)
                question_para.paragraph_format.space_before = Pt(3)
                
                # Add spacing between question and response
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6)
                
                # Response section
                response_heading = doc.add_paragraph('Analysis & Response', style='AI Section Heading')
                
                # Format the AI response with proper paragraphs
                formatted_content = format_ai_content(response)
                for content_type, content_text in formatted_content:
                    if content_type == 'heading':
                        # Sub-heading within the response
                        sub_heading = doc.add_paragraph(content_text)
                        sub_heading_run = sub_heading.runs[0]
                        sub_heading_run.font.name = 'Calibri'
                        sub_heading_run.font.size = Pt(12)
                        sub_heading_run.font.bold = True
                        sub_heading.paragraph_format.space_before = Pt(12)  # Increased from 8
                        sub_heading.paragraph_format.space_after = Pt(6)    # Increased from 4
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph(content_text, style='AI Body Text')
                        # Explicitly set spacing to ensure it's applied (python-docx style issue workaround)
                        para.paragraph_format.space_after = Pt(12)  # 12pt space_after for body text
                        para.paragraph_format.space_before = Pt(3)   # 3pt space_before for body text
                
                # Add spacing between conversations
                if i < len(ai_conversations):
                    # Add more spacing between different Q&A pairs
                    doc.add_paragraph()
                    spacer_para = doc.add_paragraph()
                    spacer_para.paragraph_format.space_after = Pt(12)
        
        elif meeting_minutes:
            # Format meeting minutes with proper structure
            minutes_heading = doc.add_paragraph('Meeting Minutes', style='AI Section Heading')
            formatted_content = format_ai_content(meeting_minutes)
            for content_type, content_text in formatted_content:
                if content_type == 'heading':
                    sub_heading = doc.add_paragraph(content_text)
                    sub_heading_run = sub_heading.runs[0]
                    sub_heading_run.font.name = 'Calibri'
                    sub_heading_run.font.size = Pt(12)
                    sub_heading_run.font.bold = True
                    sub_heading.paragraph_format.space_before = Pt(12)
                    sub_heading.paragraph_format.space_after = Pt(6)
                else:
                    para = doc.add_paragraph(content_text, style='AI Body Text')
                    # Explicitly set spacing for body text
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.space_before = Pt(3)
        
        # Footer
        if is_expeditors_user:
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Confidential - Expeditors International")
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(8)
            footer_run.font.italic = True
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create filename
        filename_base = note['title'][:30].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"AI_Analysis_{filename_base}.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
@api_router.post("/notes/{note_id}/generate-report")
async def generate_professional_report(
    note_id: str,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Generate a professional report with insights and action items"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    content = artifacts.get("transcript") or artifacts.get("text", "")
    
    if not content:
        raise HTTPException(status_code=400, detail="No content available to generate report")
    
    try:
        # Check if user is from Expeditors
        is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
        
        # Add logo header for Expeditors users
        logo_header = ""
        if is_expeditors_user:
            logo_header = """
==================================================
           EXPEDITORS INTERNATIONAL
           Professional Business Report
==================================================

"""
        
        prompt = f"""
        You are a senior executive assistant creating a professional business report{" for Expeditors International" if is_expeditors_user else ""}. Based on the following content, create a clean, well-formatted business analysis.

        Content to analyze:
        {content}

        Create a comprehensive professional report with these sections. Use clean formatting with clear headings and bullet points - NO MARKDOWN SYMBOLS like ### or **. Format as clean text with proper structure:

        EXECUTIVE SUMMARY
        Write 2-3 sentences highlighting the main points and overall situation

        KEY INSIGHTS  
        List 4-6 important findings as bullet points starting with •

        STRATEGIC RECOMMENDATIONS
        List 3-5 high-impact recommendations as bullet points starting with •

        ACTION ITEMS
        List specific, actionable next steps as bullet points starting with •
        - Include timeframes where appropriate (immediate, 1-3 months, etc.)
        - Be specific about who should be responsible

        PRIORITIES
        Categorize the action items:
        HIGH PRIORITY (next 2 weeks):
        • [list items]
        
        MEDIUM PRIORITY (next 1-3 months):
        • [list items]

        FOLLOW-UP & MONITORING
        List key metrics to track and review schedule as bullet points starting with •

        Use professional business language. Make it executive-ready. Use clear section headings in CAPS, bullet points with •, and write in complete sentences. Do NOT use markdown formatting symbols.
        """
        
        # Use enhanced AI provider with dual-provider support
        user_context = current_user.get("profile", {}) if current_user else {}
        
        analysis_result = await generate_ai_analysis(
            content=prompt,
            analysis_type="professional_report",
            user_context=user_context
        )
        
        report_content = analysis_result
        
        # Add logo header for Expeditors users
        if is_expeditors_user:
            report_content = logo_header + report_content
        
        # Store the report in artifacts
        updated_artifacts = {**artifacts, "professional_report": report_content}
        await NotesStore.set_artifacts(note_id, updated_artifacts)
        
        # Mark note as completed since report was generated
        await NotesStore.update_status(note_id, "completed")
        
        return {
            "report": report_content,
            "note_title": note["title"],
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "note_id": note_id,
            "is_expeditors": is_expeditors_user
        }
    
    except Exception as e:
        logger.error(f"Report generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Report generation temporarily unavailable")

@api_router.get("/notes/{note_id}/professional-report/export")
async def export_professional_report(
    note_id: str,
    format: str = Query("pdf", regex="^(pdf|docx|txt|rtf)$"),
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Export professional report to PDF, Word DOC, TXT, or RTF format"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    professional_report = artifacts.get("professional_report", "")
    
    if not professional_report:
        raise HTTPException(status_code=400, detail="No professional report found. Please generate a report first.")
    
    # Check if user is from Expeditors for branding
    is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
    
    # Content for export
    content_title = note['title']
    export_content = professional_report
    
    if format == "txt":
        # Plain text format
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"Professional_Report_{clean_title}.txt"
        
        return Response(
            content=export_content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "rtf":
        # RTF format
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"Professional_Report_{clean_title}.rtf"
        
        # RTF header
        rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
        
        # Convert content to RTF
        rtf_body = export_content.replace("\\", "\\\\").replace("\n", "\\par ")
        rtf_content += rtf_body + "}"
        
        return Response(
            content=rtf_content,
            media_type="application/rtf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "pdf":
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import Color, black
        from io import BytesIO
        import os
        import re
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        # Create custom styles for professional report
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'Report Title',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=18,
            spaceBefore=0,
            alignment=1,  # Center
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        # Section heading style  
        section_heading_style = ParagraphStyle(
            'Report Section Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        # Sub-heading style for content within report sections
        sub_heading_style = ParagraphStyle(
            'Report Sub Heading',
            parent=styles['Normal'],
            fontSize=12,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=black
        )
        
        # Body text style with improved spacing
        body_style = ParagraphStyle(
            'Report Body Text',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,  # Increased spacing between paragraphs
            spaceBefore=3,  # Small space before paragraphs
            leftIndent=0,
            rightIndent=0,
            fontName='Helvetica',
            leading=13.2  # 1.2 line spacing
        )
        
        # Bullet point style
        bullet_style = ParagraphStyle(
            'Report Bullet',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            spaceBefore=3,
            leftIndent=20,
            fontName='Helvetica',
            leading=13.2
        )
        
        # Function to format professional report content for PDF
        def format_report_content_pdf(content):
            if not content:
                return []
            
            # Split by double newlines to identify sections
            sections = re.split(r'\n\s*\n', content)
            formatted_elements = []
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                lines = section.split('\n')
                
                # Check if this is a section header (all caps, short line)
                if len(lines) == 1 and len(section) < 100 and section.isupper():
                    formatted_elements.append(('section_heading', section))
                elif len(lines) == 1 and len(section) < 100 and not section.endswith('.') and not section.startswith('•'):
                    # Likely a sub-heading
                    formatted_elements.append(('sub_heading', section))
                else:
                    # Process as content section
                    for line in lines:
                        line = line.strip()
                        if line:
                            if line.startswith('•'):
                                # Bullet point
                                bullet_text = line[1:].strip()
                                formatted_elements.append(('bullet', bullet_text))
                            elif line.isupper() and len(line) < 100:
                                # Section heading within content
                                formatted_elements.append(('section_heading', line))
                            else:
                                # Regular paragraph
                                formatted_elements.append(('paragraph', line))
            
            return formatted_elements
        
        story = []
        
        # Add logo if Expeditors user
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo = Image(logo_path, width=2*inch, height=0.8*inch)
                    logo.hAlign = 'CENTER'
                    story.append(logo)
                    story.append(Spacer(1, 20))
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title
        if is_expeditors_user:
            story.append(Paragraph("EXPEDITORS INTERNATIONAL", title_style))
            story.append(Paragraph("Professional Business Report", section_heading_style))
            story.append(Paragraph(f"{content_title}", body_style))
        else:
            story.append(Paragraph("Professional Business Report", title_style))
            story.append(Paragraph(f"{content_title}", section_heading_style))
        
        story.append(Spacer(1, 20))
        
        # Format the report content with proper structure
        formatted_content = format_report_content_pdf(export_content)
        for content_type, content_text in formatted_content:
            if content_type == 'section_heading':
                story.append(Paragraph(content_text, section_heading_style))
            elif content_type == 'sub_heading':
                story.append(Paragraph(content_text, sub_heading_style))
            elif content_type == 'bullet':
                story.append(Paragraph(f"• {content_text}", bullet_style))
            else:  # paragraph
                story.append(Paragraph(content_text, body_style))
        
        # Footer
        if is_expeditors_user:
            story.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer', 
                parent=styles['Normal'], 
                fontSize=8, 
                alignment=1, 
                textColor=Color(0.5, 0.5, 0.5),
                fontName='Helvetica-Oblique'
            )
            story.append(Paragraph("Confidential - Expeditors International", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        # Create filename
        filename_base = note['title'][:30].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"Professional_Report_{filename_base}.pdf"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "docx":
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from io import BytesIO
        import os
        import re
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define custom styles for better formatting
        styles = doc.styles
        
        # Main heading style
        if 'Report Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(18)
        
        # Section heading style
        if 'Report Section Heading' not in [s.name for s in styles]:
            section_style = styles.add_style('Report Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'Report Body Text' not in [s.name for s in styles]:
            body_style = styles.add_style('Report Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(12)
            body_style.paragraph_format.space_before = Pt(3)
            body_style.paragraph_format.line_spacing = 1.15
        
        # Function to format professional report content for Word
        def format_report_content_docx(content):
            if not content:
                return []
            
            # Split by double newlines to identify sections
            sections = re.split(r'\n\s*\n', content)
            formatted_elements = []
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                lines = section.split('\n')
                
                # Check if this is a section header (all caps, short line)
                if len(lines) == 1 and len(section) < 100 and section.isupper():
                    formatted_elements.append(('section_heading', section))
                elif len(lines) == 1 and len(section) < 100 and not section.endswith('.') and not section.startswith('•'):
                    # Likely a sub-heading
                    formatted_elements.append(('sub_heading', section))
                else:
                    # Process as content section
                    for line in lines:
                        line = line.strip()
                        if line:
                            if line.startswith('•'):
                                # Bullet point
                                bullet_text = line[1:].strip()
                                formatted_elements.append(('bullet', bullet_text))
                            elif line.isupper() and len(line) < 100:
                                # Section heading within content
                                formatted_elements.append(('section_heading', line))
                            else:
                                # Regular paragraph
                                formatted_elements.append(('paragraph', line))
            
            return formatted_elements
        
        # Add logo if Expeditors user
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
                    run.add_picture(logo_path, width=Inches(3))
                    doc.add_paragraph()  # Spacer
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title
        if is_expeditors_user:
            company_title = doc.add_paragraph('EXPEDITORS INTERNATIONAL', style='Report Title')
            report_title = doc.add_paragraph(f'Professional Business Report', style='Report Section Heading')
            report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_title = doc.add_paragraph(f'{content_title}', style='Report Body Text')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Explicitly set spacing for body text
            doc_title.paragraph_format.space_after = Pt(12)
            doc_title.paragraph_format.space_before = Pt(3)
        else:
            main_title = doc.add_paragraph(f'Professional Business Report', style='Report Title')
            doc_title = doc.add_paragraph(f'{content_title}', style='Report Section Heading')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a horizontal line separator
        doc.add_paragraph()
        
        # Format the report content with proper structure
        formatted_content = format_report_content_docx(export_content)
        for content_type, content_text in formatted_content:
            if content_type == 'section_heading':
                heading_para = doc.add_paragraph(content_text, style='Report Section Heading')
            elif content_type == 'sub_heading':
                # Sub-heading within the report
                sub_heading = doc.add_paragraph(content_text)
                sub_heading_run = sub_heading.runs[0]
                sub_heading_run.font.name = 'Calibri'
                sub_heading_run.font.size = Pt(12)
                sub_heading_run.font.bold = True
                sub_heading.paragraph_format.space_before = Pt(12)
                sub_heading.paragraph_format.space_after = Pt(6)
            elif content_type == 'bullet':
                bullet_para = doc.add_paragraph(f'• {content_text}', style='Report Body Text')
                # Explicitly set spacing for bullet points
                bullet_para.paragraph_format.space_after = Pt(6)
                bullet_para.paragraph_format.space_before = Pt(3)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
            else:  # paragraph
                para = doc.add_paragraph(content_text, style='Report Body Text')
                # Explicitly set spacing to ensure it's applied
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.space_before = Pt(3)
        
        # Footer
        if is_expeditors_user:
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Confidential - Expeditors International")
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(8)
            footer_run.font.italic = True
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create filename
        filename_base = note['title'][:30].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"Professional_Report_{filename_base}.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )

@api_router.post("/notes/comprehensive-batch-report")
async def generate_comprehensive_batch_report(
    request: BatchReportRequest,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Generate a comprehensive batch report including Meeting Minutes and Action Items"""
    note_ids = request.note_ids
    title = request.title
    format = request.format
    
    if not note_ids:
        raise HTTPException(status_code=400, detail="No notes provided")
    
    try:
        # Collect all content from notes
        all_transcripts = []
        all_meeting_minutes = []
        all_action_items = []
        note_titles = []
        report_date = datetime.now().strftime("%d %B %Y")
        
        for note_id in note_ids:
            note = await NotesStore.get(note_id)
            if not note:
                continue
                
            # Check user permissions
            if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
                continue
                
            artifacts = note.get("artifacts", {})
            transcript = artifacts.get("transcript") or artifacts.get("text", "")
            
            if transcript:
                note_titles.append(note["title"])
                all_transcripts.append({
                    "title": note["title"],
                    "content": transcript,
                    "created_at": note.get("created_at", "")
                })
        
        if not all_transcripts:
            raise HTTPException(status_code=400, detail="No valid content found in selected notes")
        
        # Generate Meeting Minutes for the entire batch
        # Clean speaker labels from all content first
        cleaned_transcripts = []
        for item in all_transcripts:
            content = item['content']
            # Remove speaker labels
            import re
            content = re.sub(r'Speaker \d+:\s*', '', content)
            content = re.sub(r'Speaker [A-Z]:\s*', '', content)
            content = re.sub(r'Speaker:\s*', '', content)
            cleaned_transcripts.append({
                'title': item['title'],
                'content': content,
                'created_at': item['created_at']
            })
        
        session_list = "\n\n".join([f"SESSION: {item['title']}\n{item['content']}" for item in cleaned_transcripts])
        combined_transcript = session_list
        
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            raise HTTPException(status_code=503, detail="AI service configuration error")
        
        # Generate comprehensive meeting minutes matching user's preferred format
        meeting_minutes_prompt = f"""
        Based on the following multi-session meeting transcripts, create meeting minutes following this EXACT format structure (NO markdown, NO bold formatting):

        MEETING SUMMARY

        Date: {report_date}
        Meeting Type: Multi-Session Team Meeting Summary

        EXECUTIVE SUMMARY

        This comprehensive meeting summary covers {len(all_transcripts)} sessions focused on reviewing performance, addressing operational challenges, and planning strategic initiatives. The meetings emphasized open communication, collaborative problem-solving, and clear action planning to achieve organizational objectives.

        OPENING REMARKS AND OBJECTIVES

        The sessions commenced with welcoming remarks from the chairperson, setting clear agendas and expectations for productive discussions. The primary objectives included reviewing current performance metrics, identifying areas for improvement, and establishing clear action plans for upcoming initiatives.

        KEY DISCUSSION TOPICS

        Performance Review and Analysis
        [Extract and summarize performance-related discussions from all sessions]

        Strategic Planning and Resource Allocation
        [Extract and summarize strategic planning discussions from all sessions]

        Customer Engagement and Service Delivery
        [Extract and summarize customer-related discussions from all sessions]

        Operational Efficiency and Process Improvement
        [Extract and summarize operational discussions from all sessions]

        Team Development and Communication
        [Extract and summarize team development discussions from all sessions]

        DECISIONS AND RESOLUTIONS

        [Extract key decisions made across all sessions and their collective impact]

        ACTION ITEMS AND NEXT STEPS

        [Extract action items from all sessions in narrative form, mentioning which session they came from]

        CONCLUSION

        The sessions successfully addressed key operational and strategic topics, resulting in clear action plans and renewed commitment to organizational objectives. All participants agreed to maintain open communication and collaborative approach to achieving established goals.

        TRANSCRIPTS:
        {combined_transcript}
        
        Use professional business language. Structure content logically with clear section divisions. Write in narrative form with detailed explanations. NO bullet points in main content. Focus on business outcomes and strategic initiatives.
        """
        
        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": meeting_minutes_prompt}],
                    "max_tokens": 2500,
                    "temperature": 0.2
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            response.raise_for_status()
            meeting_minutes_result = response.json()["choices"][0]["message"]["content"]
        
        # Generate consolidated action items table
        action_items_prompt = f"""
        Based on these multi-session meeting transcripts, create a clean, structured action items table:

        TRANSCRIPTS:
        {combined_transcript}

        Create a properly formatted action items table:

        CONSOLIDATED ACTION ITEMS

        No. | Action Item | Session Source | Priority | Status
        1   | [Clear, actionable description] | [Session name] | High/Medium/Low | Pending
        2   | [Clear, actionable description] | [Session name] | High/Medium/Low | Pending

        Requirements:
        - Extract ONLY clear, actionable items (not vague discussions)
        - Remove any speaker references (Speaker 1, Speaker 2, etc.)
        - Focus on concrete tasks and deliverables
        - Use professional business language
        - Assign logical priority levels based on business impact
        - Keep descriptions concise but specific
        - NO markdown formatting or special characters
        - Maximum 20 action items to keep focused
        """
        
        async with httpx.AsyncClient(timeout=120) as client:
            response = await client.post(
                "https://api.openai.com/v1/chat/completions",
                json={
                    "model": "gpt-4o-mini",
                    "messages": [{"role": "user", "content": action_items_prompt}],
                    "max_tokens": 2000,
                    "temperature": 0.2
                },
                headers={"Authorization": f"Bearer {api_key}"}
            )
            response.raise_for_status()
            action_items_result = response.json()["choices"][0]["message"]["content"]
        
        # Combine everything into comprehensive report
        if format == "ai":
            # Clean session content by removing speaker labels and unnecessary formatting
            clean_sessions = []
            for item in all_transcripts:
                # Remove speaker labels and clean up content
                content = item['content']
                
                # Remove speaker labels (Speaker 1:, Speaker 2:, Speaker A:, etc.)
                import re
                content = re.sub(r'Speaker \d+:\s*', '', content)
                content = re.sub(r'Speaker [A-Z]:\s*', '', content)
                content = re.sub(r'Speaker:\s*', '', content)
                
                # Remove excessive whitespace and empty lines
                lines = [line.strip() for line in content.split('\n') if line.strip()]
                clean_content = '\n'.join(lines)
                
                # Only include if there's substantial content
                if len(clean_content) > 100:
                    clean_sessions.append(f"\nSESSION: {item['title']}\n{'-' * 50}\n{clean_content}")
            
            session_appendix = "\n".join(clean_sessions) if clean_sessions else "Session transcripts were processed but contained limited actionable content."
            
            final_content = f"""COMPREHENSIVE MULTI-SESSION REPORT
{title}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Sessions: {len(all_transcripts)}

SOURCE NOTES:
{chr(10).join([f"• {item['title']}" for item in all_transcripts])}

---

{meeting_minutes_result}

---

{action_items_result}

---

APPENDIX: CLEANED SESSION SUMMARIES
{session_appendix}
"""
        
        else:
            # Clean format for TXT/RTF - Simple notes without speaker confusion
            clean_content = f"""{title}
Multi-Session Report
Generated: {report_date}
Sessions: {len(all_transcripts)}

"""
            # Just include the actual content without complex formatting
            for item in all_transcripts:
                clean_content += f"""
SESSION: {item['title']}
{'-' * 50}
{item['content']}

"""
            
            final_content = clean_content
        
        # Handle RTF format
        if format == "rtf":
            title_escaped = title.replace('&', '\\&')
            content_escaped = final_content.replace('&', '\\&').replace('\n', '\\par ')
            rtf_content = f"""{{\\rtf1\\ansi\\deff0 {{\\fonttbl {{\\f0 Times New Roman;}}}}
\\f0\\fs24 
\\b {title_escaped}\\b0\\par
\\par
{content_escaped}
}}"""
            final_content = rtf_content
        
        filename = f"{title.replace(' ', '_')}_Comprehensive_Report.{format if format != 'ai' else 'txt'}"
        
        # Mark all accessible notes as completed after successful export
        for note_id in note_ids:
            note = await NotesStore.get(note_id)
            if note and (not current_user or not note.get("user_id") or note.get("user_id") == current_user["id"]):
                await NotesStore.update_status(note_id, "completed")
        
        return {
            "content": final_content,
            "filename": filename,
            "note_count": len(all_transcripts),
            "title": title,
            "format": format,
            "sessions": note_titles,
            "generated_at": datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Comprehensive batch report generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Comprehensive report generation temporarily unavailable")

@api_router.post("/notes/batch-comprehensive-report")
async def generate_batch_comprehensive_report(
    request: dict,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Generate a comprehensive business report from multiple notes - same as individual professional reports"""
    note_ids = request.get("note_ids", [])
    title = request.get("title", f"Comprehensive Business Analysis - {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
    
    if not note_ids:
        raise HTTPException(status_code=400, detail="No notes provided")
    
    # Collect content from all notes
    combined_content = []
    note_titles = []
    
    for note_id in note_ids:
        note = await NotesStore.get(note_id)
        if not note:
            continue
            
        # Check access permission  
        if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
            continue
            
        artifacts = note.get("artifacts", {})
        content = artifacts.get("transcript") or artifacts.get("text", "")
        
        if content:
            note_titles.append(note['title'])
            combined_content.append(f"=== {note['title']} ===\n{content}")
    
    if not combined_content:
        raise HTTPException(status_code=400, detail="No accessible content found in the selected notes")
    
    try:
        # Use OpenAI to generate professional analysis - same as individual reports
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("WHISPER_API_KEY")
        if not api_key:
            raise HTTPException(status_code=500, detail="AI service not configured")
        
        # Check if user is from Expeditors
        is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
        
        # Combine all content
        full_content = "\n\n".join(combined_content)
        
        # Add logo header for Expeditors users
        logo_header = ""
        if is_expeditors_user:
            logo_header = """
==================================================
           EXPEDITORS INTERNATIONAL
        Comprehensive Business Analysis Report
==================================================

"""
        
        prompt = f"""
        You are a senior business analyst creating a comprehensive strategic report{" for Expeditors International" if is_expeditors_user else ""} from multiple business documents. Based on the following content from {len(note_titles)} documents, create a clean, well-formatted comprehensive business analysis.

        Documents to analyze:
        {full_content}

        CRITICAL FORMATTING REQUIREMENTS:
        - Use ONLY plain text with NO markdown symbols
        - NO asterisks (*), NO hash symbols (#), NO underscores (_)
        - Use CAPITAL LETTERS for section headings
        - Use bullet points with • symbol only
        - Write in clean, professional business language
        - Structure with clear paragraph breaks

        Create a comprehensive professional report with these sections:

        EXECUTIVE SUMMARY
        Write 2-3 sentences highlighting the main themes and strategic insights across all documents

        KEY INSIGHTS
        List 6-8 important findings that emerge from analyzing all documents together as bullet points starting with •

        STRATEGIC RECOMMENDATIONS  
        List 4-6 high-impact strategic recommendations based on the collective analysis as bullet points starting with •

        RISK ASSESSMENT
        List 3-5 potential risks or challenges identified across the documents as bullet points starting with •

        IMPLEMENTATION PRIORITIES
        List 3-5 priority actions for immediate implementation as bullet points starting with •

        CONCLUSION
        Write 2-3 sentences summarizing the overall strategic outlook and next steps

        Remember: Use ONLY plain text formatting. NO markdown symbols whatsoever.
        """
        
        # Generate AI analysis using same method as individual reports
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4000,
            temperature=0.7
        )
        
        report_content = response.choices[0].message.content
        
        # Clean any remaining markdown symbols from the AI response
        import re
        report_content = re.sub(r'\*\*\*', '', report_content)  # Remove triple asterisks
        report_content = re.sub(r'\*\*', '', report_content)    # Remove double asterisks
        report_content = re.sub(r'\*', '', report_content)      # Remove single asterisks
        report_content = re.sub(r'###', '', report_content)     # Remove triple hashes
        report_content = re.sub(r'##', '', report_content)      # Remove double hashes
        report_content = re.sub(r'#', '', report_content)       # Remove single hashes
        report_content = re.sub(r'__', '', report_content)      # Remove double underscores
        report_content = re.sub(r'_', '', report_content)       # Remove single underscores
        report_content = re.sub(r'`', '', report_content)       # Remove backticks
        
        # Clean up any extra whitespace that might result from symbol removal
        report_content = re.sub(r'\n\s*\n\s*\n', '\n\n', report_content)  # Replace multiple newlines with double newlines
        report_content = report_content.strip()
        
        # Store the report in a temporary structure (similar to how individual reports work)
        comprehensive_report_data = {
            "report": report_content,
            "title": title,
            "note_count": len(note_titles),
            "source_notes": note_titles,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "is_expeditors_user": is_expeditors_user
        }
        
        return {
            "report": report_content,
            "title": title,
            "note_count": len(note_titles),
            "source_notes": note_titles,
            "note_ids": note_ids  # Include for export reference
        }
        
    except Exception as e:
        logger.error(f"Batch comprehensive report generation error: {str(e)}")
        raise HTTPException(status_code=500, detail="Report generation temporarily unavailable")

@api_router.post("/notes/batch-comprehensive-report/export")
async def export_batch_comprehensive_report(
    request: dict,
    format: str = Query("pdf", regex="^(pdf|docx|txt|rtf)$"),
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Export batch comprehensive report - same as individual professional report export"""
    
    # Get the report data that was just generated
    report_content = request.get("report", "")
    report_title = request.get("title", "Comprehensive Business Analysis")
    
    if not report_content:
        raise HTTPException(status_code=400, detail="No report content provided")
    
    # Check if user is from Expeditors for branding
    is_expeditors_user = current_user and current_user.get("email", "").endswith("@expeditors.com")
    
    if format == "txt":
        # Plain text format - same as individual reports
        content = report_content
        if is_expeditors_user:
            content = "EXPEDITORS INTERNATIONAL\n" + "="*50 + "\n\n" + content + "\n\n" + "="*50 + "\nConfidential - Expeditors International"
        
        filename = f"Comprehensive_Report_{report_title.replace(' ', '_')[:30]}.txt"
        
        return Response(
            content=content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "rtf":
        # RTF format - same as individual reports  
        rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
        rtf_body = report_content.replace("\\", "\\\\").replace("\n", "\\par ")
        rtf_content += rtf_body + "}"
        
        filename = f"Comprehensive_Report_{report_title.replace(' ', '_')[:30]}.rtf"
        
        return Response(
            content=rtf_content,
            media_type="application/rtf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "pdf":
        # Use the SAME PDF generation as individual professional reports
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import Color, black
        from io import BytesIO
        import os
        import re
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        # Create custom styles - SAME as individual reports
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'Report Title',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=18,
            spaceBefore=0,
            alignment=1,  # Center
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        section_heading_style = ParagraphStyle(
            'Report Section Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            fontName='Helvetica-Bold',
            textColor=Color(234/255, 10/255, 42/255) if is_expeditors_user else black
        )
        
        body_style = ParagraphStyle(
            'Report Body Text',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=3,
            leftIndent=0,
            rightIndent=0,
            fontName='Helvetica',
            leading=13.2
        )
        
        story = []
        
        # Add logo if Expeditors user - SAME as individual reports
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo = Image(logo_path, width=2*inch, height=0.8*inch)
                    logo.hAlign = 'CENTER'
                    story.append(logo)
                    story.append(Spacer(1, 20))
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title - SAME structure as individual reports
        if is_expeditors_user:
            story.append(Paragraph("EXPEDITORS INTERNATIONAL", title_style))
            story.append(Paragraph("Comprehensive Business Analysis", section_heading_style))
            story.append(Paragraph(f"{report_title}", body_style))
        else:
            story.append(Paragraph("Comprehensive Business Analysis", title_style))
            story.append(Paragraph(f"{report_title}", section_heading_style))
        
        story.append(Spacer(1, 20))
        
        # Format report content - SAME as individual reports
        paragraphs = re.split(r'\n\s*\n', report_content)
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Check if it looks like a heading
                if len(paragraph) < 100 and not paragraph.endswith('.') and paragraph.isupper():
                    story.append(Paragraph(paragraph, section_heading_style))
                else:
                    story.append(Paragraph(paragraph, body_style))
        
        # Footer - SAME as individual reports
        if is_expeditors_user:
            story.append(Spacer(1, 30))
            footer_style = ParagraphStyle(
                'Footer', 
                parent=styles['Normal'], 
                fontSize=8, 
                alignment=1, 
                textColor=Color(0.5, 0.5, 0.5),
                fontName='Helvetica-Oblique'
            )
            story.append(Paragraph("Confidential - Expeditors International", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        filename = f"Comprehensive_Report_{report_title.replace(' ', '_')[:30]}.pdf"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "docx":
        # Use the SAME Word generation as individual professional reports
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from io import BytesIO
        import os
        import re
        
        doc = Document()
        
        # Set document margins - SAME as individual reports
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Define custom styles - SAME as individual reports
        styles = doc.styles
        
        if 'Report Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(18)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(18)
        
        if 'Report Section Heading' not in [s.name for s in styles]:
            section_style = styles.add_style('Report Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_font = section_style.font
            section_font.name = 'Calibri'
            section_font.size = Pt(14)
            section_font.bold = True
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
        
        if 'Report Body Text' not in [s.name for s in styles]:
            body_style = styles.add_style('Report Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
            body_style.paragraph_format.space_after = Pt(12)
            body_style.paragraph_format.space_before = Pt(3)
            body_style.paragraph_format.line_spacing = 1.15
        
        # Add logo if Expeditors user - SAME as individual reports
        if is_expeditors_user:
            logo_path = "/app/backend/expeditors-logo.png"
            if os.path.exists(logo_path):
                try:
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
                    run.add_picture(logo_path, width=Inches(3))
                    doc.add_paragraph()  # Spacer
                except Exception as e:
                    print(f"Logo loading error: {e}")
        
        # Title - SAME structure as individual reports
        if is_expeditors_user:
            company_title = doc.add_paragraph('EXPEDITORS INTERNATIONAL', style='Report Title')
            analysis_title = doc.add_paragraph(f'Comprehensive Business Analysis', style='Report Section Heading')
            analysis_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_title = doc.add_paragraph(f'{report_title}', style='Report Body Text')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_title.paragraph_format.space_after = Pt(12)
            doc_title.paragraph_format.space_before = Pt(3)
        else:
            main_title = doc.add_paragraph(f'Comprehensive Business Analysis', style='Report Title')
            doc_title = doc.add_paragraph(f'{report_title}', style='Report Section Heading')
            doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Format report content - SAME as individual reports
        paragraphs = re.split(r'\n\s*\n', report_content)
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Check if it looks like a heading
                if len(paragraph) < 100 and not paragraph.endswith('.') and paragraph.isupper():
                    heading_para = doc.add_paragraph(paragraph, style='Report Section Heading')
                else:
                    para = doc.add_paragraph(paragraph, style='Report Body Text')
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.space_before = Pt(3)
        
        # Footer - SAME as individual reports
        if is_expeditors_user:
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run("Confidential - Expeditors International")
            footer_run.font.name = 'Calibri'
            footer_run.font.size = Pt(8)
            footer_run.font.italic = True
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"Comprehensive_Report_{report_title.replace(' ', '_')[:30]}.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )

@api_router.get("/notes/{note_id}/export")
async def export_note(
    note_id: str,
    format: str = Query("txt", regex="^(txt|md|json|rtf)$"),
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Export note in various formats (txt, md, json, rtf)"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    artifacts = note.get("artifacts", {})
    
    if format == "json":
        # Convert datetime to string for JSON serialization
        export_data = {
            "id": note["id"],
            "title": note["title"],
            "kind": note["kind"],
            "created_at": note["created_at"].isoformat() if hasattr(note["created_at"], 'isoformat') else str(note["created_at"]),
            "artifacts": artifacts
        }
        # Mark note as completed since file was exported
        if current_user:  # Only update status for authenticated users
            await NotesStore.update_status(note_id, "completed")
        
        # Use note title for filename
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{clean_title}.json"
        
        return JSONResponse(
            content=export_data,
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "md":
        content = f"# {note['title']}\n\n"
        content += f"**Created:** {note['created_at']}\n"
        content += f"**Type:** {note['kind']}\n\n"
        
        if artifacts.get("transcript"):
            content += "## Transcript\n\n"
            content += artifacts["transcript"] + "\n\n"
        
        if artifacts.get("text"):
            content += "## OCR Text\n\n"
            content += artifacts["text"] + "\n\n"
        
        if artifacts.get("summary"):
            content += "## Summary\n\n"
            content += artifacts["summary"] + "\n\n"
        
        if artifacts.get("actions"):
            content += "## Action Items\n\n"
            for action in artifacts["actions"]:
                content += f"- {action}\n"
        
        # Mark note as completed since file was exported
        if current_user:  # Only update status for authenticated users
            await NotesStore.update_status(note_id, "completed")
        
        # Use note title for filename
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{clean_title}.md"
        
        return Response(
            content=content, 
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    elif format == "rtf":
        # Clean RTF format - completely raw transcript without any AI formatting
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{clean_title}.rtf"
        
        # RTF Header
        content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
        
        # Title (minimal formatting)
        content += f"\\b\\fs28 {note['title']}\\b0\\fs24\\par\\par "
        
        # Metadata  
        content += f"Created: {note['created_at']}\\par "
        content += f"Type: {note['kind']}\\par\\par "
        
        # Raw transcript content with proper paragraph formatting
        if artifacts.get("transcript"):
            raw_transcript = artifacts["transcript"]
            
            # Clean but preserve paragraph structure
            lines = raw_transcript.split('\n')
            clean_lines = []
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if line:
                    # Remove markdown formatting but keep content
                    clean_line = line.replace("**", "").replace("###", "").replace("##", "").replace("#", "")
                    clean_line = clean_line.replace("*", "").replace("_", "")
                    clean_line = clean_line.lstrip('•-*1234567890. ')
                    
                    if clean_line:
                        current_paragraph.append(clean_line)
                else:
                    # Empty line indicates paragraph break
                    if current_paragraph:
                        clean_lines.append(' '.join(current_paragraph))
                        current_paragraph = []
            
            # Add final paragraph if exists
            if current_paragraph:
                clean_lines.append(' '.join(current_paragraph))
            
            # Convert to RTF with proper paragraph breaks
            for paragraph in clean_lines:
                content += paragraph.replace("\\", "\\\\") + "\\par\\par "
        
        # Raw OCR text with paragraphs preserved
        if artifacts.get("text"):
            raw_text = artifacts["text"]
            # Clean but preserve paragraph structure
            paragraphs = raw_text.split('\n\n')  # Split on double newlines
            for paragraph in paragraphs:
                if paragraph.strip():
                    clean_paragraph = paragraph.replace("**", "").replace("###", "")
                    clean_paragraph = clean_paragraph.replace("##", "").replace("#", "")
                    clean_paragraph = clean_paragraph.replace("*", "").replace("_", "")
                    clean_paragraph = clean_paragraph.replace('\n', ' ').strip()
                    content += clean_paragraph.replace("\\", "\\\\") + "\\par\\par "
        
        # RTF Footer
        content += "}"
        
        # Mark note as completed since file was exported
        if current_user:
            await NotesStore.update_status(note_id, "completed")
        
        return Response(
            content=content, 
            media_type="application/rtf",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )
    
    else:  # txt format - completely clean raw content
        content = f"{note['title']}\n"
        content += "=" * len(note['title']) + "\n\n"
        content += f"Created: {note['created_at']}\n"
        content += f"Type: {note['kind']}\n\n"
        
        # Raw transcript with proper paragraph formatting
        if artifacts.get("transcript"):
            raw_transcript = artifacts["transcript"]
            
            # Clean but preserve natural paragraph structure
            lines = raw_transcript.split('\n')
            paragraphs = []
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if line:
                    # Remove formatting but keep content
                    clean_line = line.replace("**", "").replace("###", "").replace("##", "").replace("#", "")
                    clean_line = clean_line.replace("*", "").replace("_", "")
                    clean_line = clean_line.lstrip('•-*1234567890. ')
                    
                    # Skip AI section headers but keep the content
                    if not any(header in clean_line.upper() for header in ['ATTENDEES:', 'APOLOGIES:', 'MEETING MINUTES:', 'ACTION ITEMS:', 'KEY INSIGHTS:', 'ASSESSMENTS:', 'RISK ASSESSMENT:', 'NEXT STEPS:']):
                        if clean_line:
                            current_paragraph.append(clean_line)
                else:
                    # Empty line creates paragraph break
                    if current_paragraph:
                        paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
            
            # Add final paragraph
            if current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
            
            # Join paragraphs with double newlines for proper formatting
            content += '\n\n'.join(paragraphs) + "\n\n"
        
        # Raw OCR text with paragraph preservation
        if artifacts.get("text"):
            raw_text = artifacts["text"]
            # Clean but preserve paragraph structure
            clean_text = raw_text.replace("**", "").replace("###", "").replace("##", "").replace("#", "")
            clean_text = clean_text.replace("*", "").replace("_", "")
            content += clean_text + "\n\n"
        
        # Mark note as completed since file was exported
        if current_user:
            await NotesStore.update_status(note_id, "completed")
        
        # Use note title for filename
        clean_title = note['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
        filename = f"{clean_title}.txt"
        
        return Response(
            content=content, 
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
        )

@api_router.post("/notes/{note_id}/git-sync")
async def sync_note_to_git(
    note_id: str,
    background_tasks: BackgroundTasks,
    current_user: Optional[dict] = Depends(get_current_user_optional)
):
    """Sync note to Git repository"""
    note = await NotesStore.get(note_id)
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Check if user owns this note (if authenticated)
    if current_user and note.get("user_id") and note.get("user_id") != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized to access this note")
    
    background_tasks.add_task(enqueue_git_sync, note_id)
    return {"message": "Git sync queued"}

@api_router.get("/metrics")
async def get_metrics(
    days: int = Query(7, ge=1, le=90),
    current_user: dict = Depends(get_current_user)
) -> Dict[str, Any]:
    """Get productivity metrics for authenticated user only"""
    
    # Get stored metrics from user record first
    user_doc = await db["users"].find_one({"id": current_user["id"]})
    stored_metrics = {}
    if user_doc:
        stored_metrics = {
            "total_time_saved": user_doc.get("total_time_saved", 0),
            "notes_count": user_doc.get("notes_count", 0),
            "audio_notes_count": user_doc.get("audio_notes_count", 0),
            "photo_notes_count": user_doc.get("photo_notes_count", 0),
            "text_notes_count": user_doc.get("text_notes_count", 0),
            "avg_processing_time_minutes": user_doc.get("avg_processing_time_minutes", 0),
            "last_metrics_update": user_doc.get("last_metrics_update")
        }
    
    # Get recent notes for time-window specific metrics
    since = datetime.now(timezone.utc) - timedelta(days=days)
    query = {"created_at": {"$gte": since}, "user_id": current_user["id"]}
    
    cursor = db["notes"].find(query)
    recent_notes = await cursor.to_list(None)
    
    total_recent = len(recent_notes)
    ready_recent = sum(1 for n in recent_notes if n.get("status") in ["ready", "completed"])
    
    # Calculate success rate for recent period
    success_rate = round(ready_recent / total_recent * 100, 1) if total_recent > 0 else 100
    
    # Get latency metrics for recent notes
    latencies = [n.get("metrics", {}).get("latency_ms") for n in recent_notes if n.get("metrics", {}).get("latency_ms")]
    p95 = None
    if latencies:
        sorted_latencies = sorted(latencies)
        idx = int(0.95 * len(sorted_latencies)) - 1
        if idx >= 0:
            p95 = sorted_latencies[idx]
    
    # Calculate recent period time savings
    recent_audio = [n for n in recent_notes if n.get("kind") == "audio" and n.get("status") in ["ready", "completed"]]
    recent_photo = [n for n in recent_notes if n.get("kind") == "photo" and n.get("status") in ["ready", "completed"]]
    recent_text = [n for n in recent_notes if n.get("kind") == "text" and n.get("status") in ["ready", "completed"]]
    
    recent_time_saved = (
        len(recent_audio) * 30 +
        len(recent_photo) * 10 +
        len(recent_text) * 5
    )
    
    return {
        "window_days": days,
        
        # Recent period metrics (time-window specific)
        "recent_notes_total": total_recent,
        "recent_notes_ready": ready_recent,
        "recent_time_saved": recent_time_saved,
        "success_rate": success_rate,
        "p95_latency_ms": p95,
        
        # Overall user metrics (all-time from stored data)
        "notes_total": stored_metrics.get("notes_count", total_recent),
        "notes_audio": stored_metrics.get("audio_notes_count", len(recent_audio)),
        "notes_photo": stored_metrics.get("photo_notes_count", len(recent_photo)),
        "notes_text": stored_metrics.get("text_notes_count", len(recent_text)),
        "estimated_minutes_saved": stored_metrics.get("total_time_saved", recent_time_saved),
        "avg_processing_time_minutes": stored_metrics.get("avg_processing_time_minutes", 0),
        "last_metrics_update": stored_metrics.get("last_metrics_update"),
        
        # Metadata
        "user_authenticated": True,
        "metrics_auto_tracked": True  # Indicates metrics are automatically updated
    }

# Include the router in the main app
app.include_router(api_router)
app.include_router(upload_router, prefix="/api")  # New resumable upload API
app.include_router(transcription_router, prefix="/api")  # New transcription job API
app.include_router(webhook_router, prefix="/api")  # Phase 4: Webhook management API
app.include_router(streaming_router)  # Live Transcription Streaming API

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security headers middleware
@app.middleware("http")
async def add_security_headers(request, call_next):
    """Add security headers to all responses"""
    response = await call_next(request)
    
    # Security headers
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["Content-Security-Policy"] = "default-src 'self'"
    
    return response

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("startup")
async def startup_event():
    """Start pipeline worker and Phase 4 services on server startup"""
    from worker_manager import start_pipeline_worker
    
    # Start pipeline worker
    logger.info("🚀 Starting pipeline worker...")
    try:
        await start_pipeline_worker()
        logger.info("✅ Pipeline worker started successfully")
    except Exception as e:
        logger.error(f"❌ Failed to start pipeline worker: {e}")
    
    # Phase 4: Start production services
    logger.info("🚀 Starting Phase 4 production services...")
    
    # Start monitoring service
    try:
        await monitoring_service.start_monitoring()
        logger.info("✅ Monitoring service started")
    except Exception as e:
        logger.error(f"❌ Failed to start monitoring service: {e}")
    
    # Start webhook manager
    try:
        await webhook_manager.start()
        logger.info("✅ Webhook manager started")
    except Exception as e:
        logger.error(f"❌ Failed to start webhook manager: {e}")
    
    # Start live transcription manager
    try:
        await live_transcription_manager.initialize()
        logger.info("✅ Live transcription manager started")
    except Exception as e:
        logger.error(f"❌ Failed to start live transcription manager: {e}")
    
    logger.info("🎉 All services started successfully (including live transcription)")

@app.on_event("shutdown")
async def shutdown_db_client():
    """Shutdown database, pipeline worker, and Phase 4 services"""
    from worker_manager import stop_pipeline_worker
    
    # Stop pipeline worker
    logger.info("🛑 Stopping pipeline worker...")
    try:
        await stop_pipeline_worker()
        logger.info("✅ Pipeline worker stopped successfully")
    except Exception as e:
        logger.error(f"❌ Error stopping pipeline worker: {e}")
    
    # Phase 4: Stop production services
    logger.info("🛑 Stopping Phase 4 production services...")
    
    # Stop webhook manager
    try:
        await webhook_manager.stop()
        logger.info("✅ Webhook manager stopped")
    except Exception as e:
        logger.error(f"❌ Error stopping webhook manager: {e}")
    
    # Stop monitoring service
    try:
        await monitoring_service.stop_monitoring()
        logger.info("✅ Monitoring service stopped")
    except Exception as e:
        logger.error(f"❌ Error stopping monitoring service: {e}")
    
    client.close()
    logger.info("🎉 All services stopped successfully")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)